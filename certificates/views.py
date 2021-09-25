import os
import time
from datetime import date
import json
from io import BytesIO
import base64

from django.shortcuts import render
from django.http import HttpResponse
from django.conf import settings
from django.core.files import File
from django.core.files.base import ContentFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Mm, Pt

from fpdf import FPDF

from .forms import MakeCertificateForm
from .models import Certificate
from children.models import Child


class PDF(FPDF):
	pass


def make_certificates(request):
	if not request.user.profile.admin_access and not request.user.profile.moderator_access:
		return redirect('home')

	if request.method=='POST':
		form = MakeCertificateForm(request.POST)
		if 'generate' in request.POST:
			if form.is_valid():
				users = Child.objects.all().order_by('surname', 'name', 'name2')

				regular_url = os.path.join(settings.BASE_DIR, 'static/fonts/regular.ttf')
				bold_url = os.path.join(settings.BASE_DIR, 'static/fonts/bold.ttf')

				img_url = os.path.join(settings.BASE_DIR, 'static/img/certificate.jpg')
				sertificate_num = int(form.cleaned_data['iterable'])
				sertificate_apx = ''
				if form.cleaned_data['noniterable']:
					sertificate_apx = form.cleaned_data['divider'] + form.cleaned_data['noniterable'] 

				for user in users:
					sertificate_str_num = str(sertificate_num)+sertificate_apx
					pdf = PDF(orientation='L', unit='mm', format='A4')
					pdf.l_margin = 0.0
					pdf.r_margin = 0.0
					pdf.t_margin = 0.0
					pdf.b_margin = 0.0
					pdf.set_auto_page_break(False, margin=0)

					pdf.add_page()
					pdf.add_font('regular', '', regular_url , uni=True)
					pdf.add_font('bold', '', bold_url , uni=True)
					
					pdf.image(img_url, 0, 0, pdf.w, pdf.h)

					print(pdf.b_margin)
					pdf.set_text_color(14, 14, 14)

					pdf.set_font('regular', '', 14)
					pdf.set_xy(0.0, 81.3)
					pdf.cell(w=pdf.w, h=5.0, align='C', txt = '№' + sertificate_str_num)

					pdf.set_font('bold', '', 24)
					pdf.set_xy(0.0, 99.4)
					pdf.cell(w=pdf.w, h=5.0, align='C', txt = user.get_full_name())

					teacher = 'преподаватель ' + user.teacher.profile.get_full_name()
					if user.teacher_plus_flag:
						teacher = 'преподаватель ' + user.teacher_plus.get_full_name()
					pdf.set_font('regular', '', 20)
					pdf.set_xy(0.0, 108.8)
					pdf.cell(w=pdf.w, h=5.0, align='C', txt = teacher)


					participate = 'участвовал(а) в'
					if user.sex_valid():
						participate = 'участвовала в'
						if user.sex():
							participate = 'участвовал в'
					pdf.set_font('regular', '', 24)
					pdf.set_xy(0.0, 119.9)
					pdf.cell(w=pdf.w, h=5.0, align='C', txt = participate)

					pdf.set_font('regular', '', 20)
					pdf.set_xy(0.0, 147.7)
					pdf.cell(w=pdf.w, h=5.0, align='C', txt = 'Тема конкурса: «Воинская слава России»')


					pdf.set_font('regular', '', 16)
					pdf.set_xy(0.0, 191)
					pdf.multi_cell(w=pdf.w, h=5.0, align='C', txt = '12 декабря 2020 года')

					pdf.output(os.path.join(settings.MEDIA_ROOT, 'test.pdf'), 'F')
					file  = open(os.path.join(settings.MEDIA_ROOT, 'test.pdf'), 'rb')
					djangofile = File(file)

					new_certificate = Certificate.objects.create(child = user, certificate_num = sertificate_str_num)
					new_certificate.certificate_file.save((user.get_file_name())+'.pdf', djangofile)

					file.close()

					sertificate_num += 1
			else:
				form = MakeCertificateForm(request.POST)
				certificates = Certificate.objects.all()
				args = {
					'form': form,
					'certificates': certificates
				}
				return render(request, 'certificates/certificates.html', args)

		elif 'delete' in request.POST:
				certificates = Certificate.objects.all().delete()
		else:
			certificates = Certificate.objects.all()

			dte = date.today()
			document = Document()
			section = document.sections[-1]
			new_width, new_height = section.page_height, section.page_width
			section.orientation = WD_ORIENT.PORTRAIT
			section.page_width = Mm(297)
			section.page_height = Mm(210)
			section.left_margin = Mm(30)
			section.right_margin = Mm(10)
			section.top_margin = Mm(10)
			section.bottom_margin = Mm(10)
			section.header_distance = Mm(10)
			section.footer_distance = Mm(10)

			style = document.styles['Normal']
			font = style.font
			font.name = 'Times New Roman'
			font.size = Pt(12)


			document.add_paragraph('Список сертификатов').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			p = document.add_paragraph()
			p.add_run(dte.strftime('%d.%b.%Y')).italic = True
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

			table = document.add_table(rows=1, cols=5)
			table.allow_autifit = False
			table.style = 'TableGrid'
			table.columns[0].width = Mm(10)
			table.columns[1].width = Mm(70)
			table.columns[2].width = Mm(70)
			table.columns[3].width = Mm(70)
			table.columns[4].width = Mm(37)

			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = '№'
			hdr_cells[0].width = Mm(10)
			hdr_cells[1].text = 'Конкурсант'
			hdr_cells[1].width = Mm(70)
			hdr_cells[2].text = 'Преподаватель'
			hdr_cells[2].width = Mm(70)
			hdr_cells[3].text = 'Учреждение'
			hdr_cells[3].width = Mm(70)
			hdr_cells[4].text = 'Номер сертификата'
			hdr_cells[4].width = Mm(37)


			for certificate in certificates:
				row_cells = table.add_row().cells
				row_cells[0].text = str(certificate.child.id)
				row_cells[0].width = Mm(10)
				row_cells[1].text = certificate.child.get_full_name()
				row_cells[1].width = Mm(70)
				if certificate.child.teacher_plus:
					row_cells[2].text = certificate.child.teacher_plus.get_full_name()
				else:
					row_cells[2].text =certificate.child.teacher.profile.get_full_name()
				row_cells[2].width = Mm(70)
				row_cells[3].text = certificate.child.teacher.profile.institution
				row_cells[3].width = Mm(70)
				row_cells[4].text = certificate.certificate_num
				row_cells[4].width = Mm(37)


			response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
			response['Content-Disposition'] = 'attachment; filename=Certificates (' + dte.strftime('%d-%b-%Y') + ').docx'
			document.save(response)

			return response



	form = MakeCertificateForm()
	certificates = Certificate.objects.all()
	args = {
		'form': form,
		'certificates': certificates
	}
	return render(request, 'certificates/certificates.html', args)


# --------------------------------
#           Для ajax'а
# --------------------------------
def ajax_generate_certificates(request):
	users = Child.objects.all().order_by('surname', 'name', 'name2')

	
	font_url = os.path.join(settings.BASE_DIR, 'static/fonts/chekhovskoy.ttf')
	img_url = os.path.join(settings.BASE_DIR, 'static/img/sertificate.jpg')
	sertificate_num = 33

	for user in users:
		sertificate_str_num = str(sertificate_num)+'-20'
		pdf = PDF(orientation='L', unit='mm', format='A4')
		pdf.add_page()
		pdf.add_font('Chehkovskoy', '', font_url , uni=True)
		
		if user.speaker:
			pdf.image(img_speaker_url, 0, 0, pdf.w, pdf.h)
		else:
			pdf.image(img_member_url, 0, 0, pdf.w, pdf.h)
		
		pdf.set_font('Chehkovskoy', '', 18)
		pdf.set_text_color(0, 0, 0)
		pdf.set_xy(0.0, 110.0)
		pdf.cell(w=297.0, h=5.0, align='C', txt = '№' + sertificate_str_num)
		pdf.set_xy(0.0, 130.0)
		pdf.cell(w=297.0, h=5.0, align='C', txt = user.get_full_name())

		pdf.output(os.path.join(settings.MEDIA_ROOT, 'test.pdf'), 'F')
		file  = open(os.path.join(settings.MEDIA_ROOT, 'test.pdf'), 'rb')
		djangofile = File(file)

		user.certificate_file.save((user.get_file_name())+'.pdf', djangofile)
		user.certificate_num = sertificate_str_num
		user.save()

		file.close()

		sertificate_num += 1
	sertificate_num -= 1

	return HttpResponse(json.dumps({'last_num': str(sertificate_num) + '-20'}))

