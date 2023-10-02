import os
import datetime
import json
import locale

from datetime import date

from django.conf import settings
from django.core.files.storage import default_storage

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Pt

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.models import User

from .forms import MovieUploadForm, MovieEditForm
from .models import Movie, CoMovie
from profileuser.models import Profile, CoProfile


@login_required(login_url='/login/')
def view_movies(request):
	movies = Movie.objects.filter(author__main_user=request.user)
	comovies_list = {}
	for movie in movies:
		comovies_list[movie.pk] = CoMovie.objects.filter(movie = movie)

	if request.method=='POST':
		if 'addpict' in request.POST:
			return redirect('movies:load_movie')

		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()

		document = Document()

		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.PORTRAIT
		section.page_width = Mm(210)
		section.page_height = Mm(297)
		section.left_margin = Mm(20)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(14)

		movie_cnt = movies.count()
		for movie in movies:

			p = document.add_paragraph()
			p.add_run('ЗАЯВКА').bold = True
			p.paragraph_format.space_after = 0
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

			p = document.add_paragraph()
			p.add_run('на участие во Всероссийском фестивале-конкурсе').bold = True
			p.paragraph_format.space_after = 0
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

			p = document.add_paragraph()
			p.add_run('народного творчества «Гавриловские гуляния»').bold = True
			p.paragraph_format.space_after = 0
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

			p = document.add_paragraph()
			p.add_run('(концертная программа)').bold = True
			p.paragraph_format.space_after = 0
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

			p = document.add_paragraph()
			p.add_run(dte.strftime('%d %B %Y') + ' года').italic = True
			p.paragraph_format.space_after = 0
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

			document.add_paragraph().paragraph_format.space_after = 0


			p = document.add_paragraph()
			p.add_run('1. ФИО участника или название коллектива (полностью, без сокращений): ' + movie.author.get_full_name())
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('2. Дата рождения: ' + movie.age)
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('3. Название направляющей организации (в соответствии с ЕГРЮЛ полное и сокращенное название): ' + request.user.profile.get_institute_full())
			if not request.user.profile.less_institution:
				p.add_run(' (' + request.user.profile.get_institute() + ')')
			p.paragraph_format.space_after = 0


			p = document.add_paragraph()
			p.add_run('4. Почтовый адрес и индекс направляющей организации: ')
			if not request.user.profile.less_institution:
				p.add_run(request.user.profile.adress)
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('5. ФИО ответственного лица: ' + request.user.profile.get_full_name())
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('6. Телефон: ' + request.user.profile.phone)
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('7. E-mail: ' + request.user.email)
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('8. ФИО преподавателя, руководителя коллектива, концертмейстера (полностью, без сокращений): ')
			comovies_flag=False
			for comovies in comovies_list[movie.pk]:
				if comovies_flag:
					p.add_run(', ')
				else:
					p.add_run(' ')
				comovies_flag = True
				p.add_run(comovies.coauthor.short_profile_type() + ' ' + comovies.coauthor.get_full_name())
			p.paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('9. Программа конкурсного выступления:')
			p.paragraph_format.space_after = 0

			document.add_paragraph().paragraph_format.space_after = 0

			table = document.add_table(rows=1, cols=2)
			table.allow_autifit = False
			table.style = 'Table Grid'
			table.columns[0].width = Mm(110)
			table.columns[1].width = Mm(70)



			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Название, автор'
			hdr_cells[0].paragraphs[0].runs[0].font.bold = True
			hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[0].width = Mm(110)
			hdr_cells[1].text = 'Регион'
			hdr_cells[1].paragraphs[0].runs[0].font.bold = True
			hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[1].width = Mm(70)



			row_cells = table.add_row().cells

			row_cells[0].text = movie.name_1
			if movie.composer_1:
				row_cells[0].text += ', муз. ' + movie.composer_1
			if movie.poet_1:
				row_cells[0].text += ', сл. ' + movie.poet_1
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(110)

			row_cells[1].text = str(movie.region_1)
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(70)


			row_cells = table.add_row().cells

			row_cells[0].text = movie.name_2
			if movie.composer_2:
				row_cells[0].text += ', муз. ' + movie.composer_2
			if movie.poet_2:
				row_cells[0].text += ', сл. ' + movie.poet_2
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(110)

			row_cells[1].text = str(movie.region_2)
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(70)

			document.add_paragraph().paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('10. Технические требования для выступления:')
			p.add_run(movie.descritpion)
			p.paragraph_format.space_after = 0

			document.add_paragraph().paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('Подпись ответственного лица: _________________________________________')
			p.paragraph_format.space_after = 0

			document.add_paragraph().paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('Дата: «_______» _______________ ' + str(dte.year) + ' г.')
			p.paragraph_format.space_after = 0

			document.add_paragraph().paragraph_format.space_after = 0

			p = document.add_paragraph()
			p.add_run('Печать')
			p.paragraph_format.space_after = 0
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

			movie_cnt -= 1
			if movie_cnt:
				document.add_page_break()

		file_name = 'Statements-vocal'
		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=' + file_name +' (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response

	args = {
		'movies': movies,
		'comovies_list': comovies_list,
		'memeber_flag': True
	}
	return render(request, 'movies/view_movies.html', args)


@login_required(login_url='/login/')
def view_movies_admin(request, pk = None):
	user = User.objects.get(pk = pk)
	user_name = user.profile.get_institute_zip()

	movies = Movie.objects.filter(author__main_user=user)
	comovies_list = {}
	for movie in movies:
		comovies_list[movie.pk] = CoMovie.objects.filter(movie = movie)

	if request.method=='POST':
		if 'addpict' in request.POST:
			return redirect('movies:load_movie_admin', pk = pk)


	args = {
		'movies': movies,
		'comovies_list': comovies_list,
		'user_name': user_name,
		'admin_flag': True
	}
	return render(request, 'movies/view_movies.html', args)


@login_required(login_url='/login/')
def load_movie(request):
	if not request.user.profile.member_access:
		return redirect('home')

	author = request.user
	coprofiles = CoProfile.objects.filter(main_user = author, profile_type__in = ['1', '2', '3', '4', '5'])
	ages_pk = list(CoProfile.objects.filter(main_user = request.user, profile_type = '0', coprofile_type = '1').values_list('pk', flat=True))

	if request.method=='POST':
		form_movie = MovieUploadForm(request.POST, author = author, label_suffix='')

		if form_movie.is_valid():
			new_movie = form_movie.save(commit=False)

			if new_movie.participation == '2':
				if new_movie.file_1:
					if 'youtu' in new_movie.file_1:
						url = new_movie.file_1.split('=')
						if len(url)>1:
							url = url[1].split('&')
							new_movie.file_1 = 'https://www.youtube.com/embed/' + url[0]
						else:
							url = new_movie.file_1.split('/')
							new_movie.file_1 = 'https://www.youtube.com/embed/' + url[len(url)-1]
						new_movie.youtube_flag_1 = True
				
				if new_movie.file_2:
					if 'youtu' in new_movie.file_2:
						url = new_movie.file_2.split('=')
						if len(url)>1:
							url = url[1].split('&')
							new_movie.file_2 = 'https://www.youtube.com/embed/' + url[0]
						else:
							url = new_movie.file_2.split('/')
							new_movie.file_2 = 'https://www.youtube.com/embed/' + url[len(url)-1]
						new_movie.youtube_flag_2 = True

			new_movie.save()	

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoMovie.objects.filter(movie = new_movie, coauthor = coprofile):
						CoMovie.objects.create(movie = new_movie, coauthor = coprofile)
				else:
					CoMovie.objects.filter(movie = new_movie, coauthor = coprofile).delete()


			return redirect('movies:view_movies')

		args = {
			'form': form_movie,
			'coprofiles': coprofiles,
			'author': author,
			'ages_pk': ages_pk
		}
		return render(request, 'movies/load_movie.html', args)	
	
	form_movie = MovieUploadForm(author = author, label_suffix='')


	args = {
		'form': form_movie,
		'coprofiles': coprofiles,
		'author': author,
		'ages_pk': ages_pk
	}
	return render(request, 'movies/load_movie.html', args)


@login_required(login_url='/login/')
def load_movie_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	author = User.objects.get(pk = pk)
	user_name = author.profile.get_institute_zip()

	coprofiles = CoProfile.objects.filter(main_user = author, profile_type__in = ['1', '2', '3', '4', '5'])
	ages_pk = list(CoProfile.objects.filter(main_user = author, profile_type = '0', coprofile_type = '1').values_list('pk', flat=True))
	if request.method=='POST':
		form_movie = MovieUploadForm(request.POST, author = author, label_suffix='')

		if form_movie.is_valid():
			new_movie = form_movie.save(commit=False)

			if new_movie.participation == '2':
				if new_movie.file_1:
					if 'youtu' in new_movie.file_1:
						url = new_movie.file_1.split('=')
						if len(url)>1:
							url = url[1].split('&')
							new_movie.file_1 = 'https://www.youtube.com/embed/' + url[0]
						else:
							url = new_movie.file_1.split('/')
							new_movie.file_1 = 'https://www.youtube.com/embed/' + url[len(url)-1]
						new_movie.youtube_flag_1 = True

				if new_movie.file_2:
					if 'youtu' in new_movie.file_2:
						url = new_movie.file_2.split('=')
						if len(url)>1:
							url = url[1].split('&')
							new_movie.file_2 = 'https://www.youtube.com/embed/' + url[0]
						else:
							url = new_movie.file_2.split('/')
							new_movie.file_2 = 'https://www.youtube.com/embed/' + url[len(url)-1]
						new_movie.youtube_flag_2 = True

			new_movie.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoMovie.objects.filter(movie = new_movie, coauthor = coprofile):
						CoMovie.objects.create(movie = new_movie, coauthor = coprofile)
				else:
					CoMovie.objects.filter(movie = new_movie, coauthor = coprofile).delete()


			return redirect('movies:view_movies_admin', pk = pk)

		args = {
			'form': form_movie,
			'coprofiles': coprofiles,
			'author': author,
			'ages_pk': ages_pk,
			'user_name': user_name
		}
		return render(request, 'movies/load_movie.html', args)

	form_movie = MovieUploadForm(author = author, label_suffix='')


	args = {
		'form': form_movie,
		'coprofiles': coprofiles,
		'author': author,
		'ages_pk': ages_pk,
		'user_name': user_name
	}
	return render(request, 'movies/load_movie.html', args)


@login_required(login_url='/login/')
def edit_movie(request, pk):
	if not request.user.profile.member_access:
		return redirect('home')
		
	movie = Movie.objects.get(pk=pk)
	author = request.user
	participation = movie.participation

	comovies_list = {}
	ages_pk = list(CoProfile.objects.filter(main_user = request.user, profile_type = '0', coprofile_type = '1').values_list('pk', flat=True))

	if movie.author.main_user != author:
		return redirect('home')

	coprofiles = CoProfile.objects.filter(main_user = author, profile_type__in = ['1', '2', '3', '4', '5'])

	if request.method=='POST':
		form_movie = MovieEditForm(request.POST, instance=movie, author = author, label_suffix='')
		if form_movie.is_valid():
			new_movie = form_movie.save(commit=False)
			if new_movie.file_1:
				new_movie.youtube_flag_1 = False
				if 'youtu' in new_movie.file_1:
					url = new_movie.file_1.split('=')
					if len(url)>1:
						url = url[1].split('&')
						new_movie.file_1 = 'https://www.youtube.com/embed/' + url[0]
					else:
						url = new_movie.file_1.split('/')
						new_movie.file_1 = 'https://www.youtube.com/embed/' + url[len(url)-1]
					new_movie.youtube_flag_1 = True

			if new_movie.file_2:
				new_movie.youtube_flag_2 = False
				if 'youtu' in new_movie.file_2:
					url = new_movie.file_2.split('=')
					if len(url)>1:
						url = url[1].split('&')
						new_movie.file_2 = 'https://www.youtube.com/embed/' + url[0]
					else:
						url = new_movie.file_2.split('/')
						new_movie.file_2 = 'https://www.youtube.com/embed/' + url[len(url)-1]
					new_movie.youtube_flag_2 = True

			new_movie.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoMovie.objects.filter(movie = movie, coauthor = coprofile):
						CoMovie.objects.create(movie = movie, coauthor = coprofile)
				else:
					CoMovie.objects.filter(movie = movie, coauthor = coprofile).delete()

			return redirect('movies:view_movies')

		args = {
			'form': form_movie,
			'movie': movie,
			'coprofiles': coprofiles,
			'comovies_list': comovies_list,
			'ages_pk': ages_pk,
			'participation': participation
		}
		return render(request, 'movies/edit_movie.html', args)	
	
	form_movie = MovieEditForm(instance=movie, author = author, label_suffix='')
	comovies = list(CoMovie.objects.filter(movie = movie).values_list('coauthor', flat=True))
	for coprofile in coprofiles:
		comovies_list[coprofile.pk] = coprofile.pk in comovies


	args = {
		'form': form_movie,
		'movie': movie,
		'coprofiles': coprofiles,
		'comovies_list': comovies_list,
		'ages_pk': ages_pk,
		'participation': participation
	}
	return render(request, 'movies/edit_movie.html', args)


@login_required(login_url='/login/')
def edit_movie_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')
		
	movie = Movie.objects.get(pk=pk)
	participation = movie.participation
	author = movie.author.main_user
	user_name = author.profile.get_institute_zip()

	comovies_list = {}
	ages_pk = list(CoProfile.objects.filter(main_user = author, profile_type = '0', coprofile_type = '1').values_list('pk', flat=True))

	if movie.author.main_user != author:
		return redirect('home')

	coprofiles = CoProfile.objects.filter(main_user = author, profile_type__in = ['1', '2', '3', '4', '5'])

	if request.method=='POST':
		form_movie = MovieEditForm(request.POST, instance=movie, author = author, label_suffix='')
		if form_movie.is_valid():
			new_movie = form_movie.save(commit=False)
			if new_movie.file_1:
				new_movie.youtube_flag_1 = False
				if 'youtu' in new_movie.file_1:
					url = new_movie.file_1.split('=')
					if len(url)>1:
						url = url[1].split('&')
						new_movie.file_1 = 'https://www.youtube.com/embed/' + url[0]
					else:
						url = new_movie.file_1.split('/')
						new_movie.file_1 = 'https://www.youtube.com/embed/' + url[len(url)-1]
					new_movie.youtube_flag_1 = True

			if new_movie.file_2:
				new_movie.youtube_flag_2 = False
				if 'youtu' in new_movie.file_2:
					url = new_movie.file_2.split('=')
					if len(url)>1:
						url = url[1].split('&')
						new_movie.file_2 = 'https://www.youtube.com/embed/' + url[0]
					else:
						url = new_movie.file_2.split('/')
						new_movie.file_2 = 'https://www.youtube.com/embed/' + url[len(url)-1]
					new_movie.youtube_flag_2 = True

			new_movie.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoMovie.objects.filter(movie = movie, coauthor = coprofile):
						CoMovie.objects.create(movie = movie, coauthor = coprofile)
				else:
					CoMovie.objects.filter(movie = movie, coauthor = coprofile).delete()

			return redirect('movies:view_movies_admin', pk = author.pk)

		args = {
			'form': form_movie,
			'movie': movie,
			'coprofiles': coprofiles,
			'comovies_list': comovies_list,
			'ages_pk': ages_pk,
			'participation': participation,
			'user_name': user_name
		}
		return render(request, 'movies/edit_movie.html', args)

	form_movie = MovieEditForm(instance=movie, author = author, label_suffix='')
	comovies = list(CoMovie.objects.filter(movie = movie).values_list('coauthor', flat=True))
	for coprofile in coprofiles:
		comovies_list[coprofile.pk] = coprofile.pk in comovies


	args = {
		'form': form_movie,
		'movie': movie,
		'coprofiles': coprofiles,
		'comovies_list': comovies_list,
		'ages_pk': ages_pk,
		'participation': participation,
		'user_name': user_name
	}
	return render(request, 'movies/edit_movie.html', args)



# --------------------------------
#           Для ajax'а
# --------------------------------
@login_required(login_url='/login/')
def ajax_del_movie(request):
	movie_pk = request.GET['movie']
	movie = Movie.objects.get(pk=movie_pk)

	if not request.user.profile.admin_access and movie.author.main_user != request.user:
		return HttpResponse(False)

	movie.delete()

	return HttpResponse(True)


@login_required(login_url='/login/')
def ajax_change_scene_movie(request):
	data_list = json.loads(request.GET['data_list'])


	
	for key, val in data_list.items():
		movie = Movie.objects.get(pk=key)
		movie.scene_num = int(val)
		movie.save()

	return HttpResponse(True)
