import os
import io
import time
import locale
import json

from zipfile import ZipFile
from pytils import translit

from django.db.models import Q
from datetime import date
from io import BytesIO
from django.core.files import File
from django.core.files.base import ContentFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Pt

from fpdf import FPDF

from django.conf import settings
from django.templatetags.static import static
from django.contrib.staticfiles.storage import staticfiles_storage
from django.core.files.storage import default_storage

from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.urls import reverse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import login, authenticate, logout, update_session_auth_hash
from django.contrib.auth.models import User

from django.contrib.sites.shortcuts import get_current_site
from django.utils.encoding import force_bytes, force_text
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.template.loader import render_to_string
from django.core.mail import EmailMessage, send_mail
from .tokens import accaunt_activation_token

from profileuser.models import Profile, CoProfile

from pictures.models import Picture, CoPicturee
from movies.models import Movie, CoMovie
from invoices.models import Invoice
from protocols.models import Protocol
from statements.models import Statement
from personals.models import Personal
from nominations.models import ArtNomination, VocalNomination

from .forms import UserLoginForm, UserRegistrationForm, ChangePasswordForm, CustomPasswordResetForm, CustomSetPasswordForm
from .forms import SetJuriForm, ChangeJuriForm, NewJuriForm

class PDF(FPDF):
	pass

def home_view(request):

	args = {
	}
	return render(request, 'index.html', args)


def policy_view(request):
	return render(request, 'policy.html')


def login_view(request):
	form = UserLoginForm(request.POST or None)
	next_ = request.GET.get('next')
	modal = False

	if form.is_valid():
		username = request.POST.get('email').lower()
		password = request.POST.get('password')
		user = authenticate(username = username.strip(), password = password.strip())
		
		login(request, user)
		next_post = request.POST.get('next')
		redirect_path = next_ or next_post or '/'


		return redirect(redirect_path)

	args = {
		'form': form
	}
	return render(request, 'login.html', args)


def logout_view(request):
	logout(request)
	return redirect('home')


def register_view(request):
	if request.method=='POST':
		user_form = UserRegistrationForm(request.POST)
		if user_form.is_valid():
			new_user = user_form.save(commit=False)
			new_user.username = new_user.email
			new_user.is_active = False
			new_user.set_password(user_form.cleaned_data['password'])
			new_user.save()

			new_user.profile.less_institution = user_form.cleaned_data['less_institution']
			new_user.profile.category = user_form.cleaned_data['category']
			new_user.profile.surname = user_form.cleaned_data['surname']
			new_user.profile.name = user_form.cleaned_data['name']
			new_user.profile.name2 = user_form.cleaned_data['name2']
			new_user.profile.profile_type = user_form.cleaned_data['profile_type']
			new_user.profile.phone = user_form.cleaned_data['phone']
			new_user.profile.institution = user_form.cleaned_data['institution']
			new_user.profile.institution_shot = user_form.cleaned_data['institution_shot']
			new_user.profile.adress = user_form.cleaned_data['adress']
			
			new_user.profile.save()

			if user_form.cleaned_data['add_team']:
				CoProfile.objects.create(
					main_user = new_user,
					coprofile_type = '2',
					team = user_form.cleaned_data['team']
				)

			current_site = get_current_site(request)
			protocol = 'http'
			if request.is_secure():
				protocol = 'https'

			mail_subject = 'Активация аккаунта'
			to_email = new_user.email
			uid = urlsafe_base64_encode(force_bytes(new_user.pk))

			token = accaunt_activation_token.make_token(new_user)

			message = render_to_string('acc_active_email.html', {'user': new_user, 'domain': current_site.domain,\
				'uid': uid, \
				'token': token,\
				'protocol': protocol,\
				'email': to_email})

			message_html = render_to_string('acc_active_email_html.html', {'user': new_user, 'domain': current_site.domain,\
				'uid': uid, \
				'token': token,\
				'protocol': protocol,\
				'email': to_email})

			send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

			return render(request, 'confirm_email.html')

		return render(request, 'register.html', {'form': user_form})

	user_form = UserRegistrationForm()
	return render(request, 'register.html', {'form': user_form})


def activate(request, uidb64, token):
	try:
		uid = force_text(urlsafe_base64_decode(uidb64))
		user = User.objects.get(pk=uid)
	except(TypeError, ValueError, OverflowError, User.DoesNotExist):
		user = None

	if user and accaunt_activation_token.check_token(user, token):
		user.is_active = True
		user.save()

		login(request, user)
		return redirect('home')

	return render(request, 'failed_email.html')


def change_password(request):
	username = request.user.username

	if request.method == 'POST':
		form = ChangePasswordForm(request.POST, username=username)

		if form.is_valid():
			user = request.user
			user.set_password(form.cleaned_data['newpassword1'])
			user.save()
			update_session_auth_hash(request, user)

			return redirect('profiles:view_edit_profile')

		args = {
			'form': form
		}
		return render(request, 'change_password.html', args)

	form = ChangePasswordForm(username=username)

	args = {
		'form': form
	}
	return render(request, 'change_password.html', args)


@login_required(login_url='/login/')
def statistic_contestant_add_view(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	nomination_picture_list = dict(ArtNomination.objects.all().values_list('id', 'name'))
	nomination_move_list = dict(VocalNomination.objects.all().values_list('id', 'name'))


	near_pictures = Picture.objects.filter(participation = '1').order_by('author__coprofile_type', 'author__team','author__surname', 'author__name', 'author__name2')
	far_pictures = Picture.objects.filter(participation = '2').order_by('author__coprofile_type', 'author__team','author__surname', 'author__name', 'author__name2')
	near_moves = Movie.objects.filter(participation = '1').order_by('author__coprofile_type', 'author__team','author__surname', 'author__name', 'author__name2')
	far_moves = Movie.objects.filter(participation = '2').order_by('author__coprofile_type', 'author__team','author__surname', 'author__name', 'author__name2')

	co_teachers = {}
	co_teachers_vocal = {}
	for near_picture in near_pictures:
		co_teachers_pk = list(CoPicturee.objects.filter(picture = near_picture).values_list('coauthor', flat=True))
		co_teachers[near_picture.pk] = CoProfile.objects.filter(pk__in = co_teachers_pk)

	for far_picture in far_pictures:
		co_teachers_pk = list(CoPicturee.objects.filter(picture = far_picture).values_list('coauthor', flat=True))
		co_teachers[far_picture.pk] = CoProfile.objects.filter(pk__in = co_teachers_pk)

	for near_move in near_moves:
		co_teachers_pk = list(CoMovie.objects.filter(movie = near_move).values_list('coauthor', flat=True))
		co_teachers_vocal[near_move.pk] = CoProfile.objects.filter(pk__in = co_teachers_pk)

	for far_move in far_moves:
		co_teachers_pk = list(CoMovie.objects.filter(movie = far_move).values_list('coauthor', flat=True))
		co_teachers_vocal[far_move.pk] = CoProfile.objects.filter(pk__in = co_teachers_pk)


	near_pictures_nominations = {}
	for key,val in nomination_picture_list.items():
		near_pictures_nominations[key] = near_pictures.filter(nomination = key)

	far_pictures_nominations = {}
	for key,val in nomination_picture_list.items():
		far_pictures_nominations[key] = far_pictures.filter(nomination = key)

	near_moves_nominations = {}
	for key,val in nomination_move_list.items():
		near_moves_nominations[key] = near_moves.filter(nomination = key)

	far_moves_nominations = {}
	for key,val in nomination_move_list.items():
		far_moves_nominations[key] = far_moves.filter(nomination = key)



	if request.POST:
		if 'type_1' in request.POST or 'type_2' in request.POST or 'type_3' in request.POST or 'type_4' in request.POST:
			locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

			dte = date.today()
			document = Document()
			section = document.sections[-1]
			new_width, new_height = section.page_height, section.page_width
			section.orientation = WD_ORIENT.LANDSCAPE
			section.page_width = Mm(297)
			section.page_height = Mm(210)
			section.left_margin = Mm(30)
			section.right_margin = Mm(10)
			section.top_margin = Mm(10)
			section.bottom_margin = Mm(10)
			section.header_distance = Mm(10)
			section.footer_distance = Mm(10)

			style = document.styles['Normal']
			font = style.font
			font.name = 'Times New Roman'
			font.size = Pt(12)

			if 'type_1' in request.POST:
				document.add_paragraph('ДПИ (очное участие)').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				array_mambers_list = near_pictures
				nomination_list = nomination_picture_list
				co_teachers_docx = co_teachers
				file_name = 'dpi_near'
			elif 'type_2' in request.POST:
				document.add_paragraph('ДПИ (заочное участие)').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				array_mambers_list = far_pictures
				nomination_list = nomination_picture_list
				co_teachers_docx = co_teachers
				file_name = 'dpi_far'
			elif 'type_3' in request.POST:
				document.add_paragraph('Вокал (очное участие)').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				array_mambers_list = near_moves
				nomination_list = nomination_move_list
				co_teachers_docx = co_teachers_vocal
				file_name = 'vocal_near'
			else:
				document.add_paragraph('Вокал (заочное участие)').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				array_mambers_list = far_moves
				nomination_list = nomination_move_list
				co_teachers_docx = co_teachers_vocal
				file_name = 'vocal_far'

			p = document.add_paragraph()
			p.add_run(dte.strftime('%d %B %Y')).italic = True
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

			for category in ['1', '2', '3', '4', '5', '6']:
				if array_mambers_list.filter(author__main_user__profile__category= category).count():
					p = document.add_paragraph()
					p.add_run(dict(Profile.CATEGORY_TYPES)[category]).bold = True
					p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER


					for key, val in nomination_list.items():
						array_mambers = array_mambers_list.filter(author__main_user__profile__category= category, nomination=key)

						if array_mambers:
							p = document.add_paragraph()
							p.add_run(val).italic = True


							if 'type_3' in request.POST:
								width_list = [10, 47, 50, 50, 50, 50]
								table = document.add_table(rows=1, cols=6)
								table.allow_autifit = False
								table.style = 'Table Grid'
								table.columns[0].width = Mm(width_list[0])
								table.columns[1].width = Mm(width_list[1])
								table.columns[2].width = Mm(width_list[2])
								table.columns[3].width = Mm(width_list[3])
								table.columns[4].width = Mm(width_list[4])
								table.columns[5].width = Mm(width_list[5])
							else:
								width_list = [10, 77, 50, 50, 70]
								table = document.add_table(rows=1, cols=5)
								table.allow_autifit = False
								table.style = 'Table Grid'
								table.columns[0].width = Mm(width_list[0])
								table.columns[1].width = Mm(width_list[1])
								table.columns[2].width = Mm(width_list[2])
								table.columns[3].width = Mm(width_list[3])
								table.columns[4].width = Mm(width_list[4])



							hdr_cells = table.rows[0].cells
							hdr_cells[0].text = '№'
							hdr_cells[0].paragraphs[0].runs[0].font.bold = True
							hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							hdr_cells[0].width = Mm(width_list[0])
							hdr_cells[1].text = 'Конкурсант'
							hdr_cells[1].paragraphs[0].runs[0].font.bold = True
							hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							hdr_cells[1].width = Mm(width_list[1])
							hdr_cells[2].text = 'Название'
							hdr_cells[2].paragraphs[0].runs[0].font.bold = True
							hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							hdr_cells[2].width = Mm(width_list[2])
							hdr_cells[3].text = 'Преподаватель(и)'
							hdr_cells[3].paragraphs[0].runs[0].font.bold = True
							hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							hdr_cells[3].width = Mm(width_list[3])
							hdr_cells[4].text = 'Учреждение'
							hdr_cells[4].paragraphs[0].runs[0].font.bold = True
							hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							hdr_cells[4].width = Mm(width_list[4])
							if 'type_3' in request.POST:
								hdr_cells[5].text = 'Тех.требования'
								hdr_cells[5].paragraphs[0].runs[0].font.bold = True
								hdr_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
								hdr_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
								hdr_cells[5].width = Mm(width_list[5])

							cnt = 0
							for member in array_mambers:
								cnt += 1
								row_cells = table.add_row().cells
								row_cells[0].text = str(cnt)
								row_cells[0].paragraphs[0].runs[0].font.bold = True
								row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
								row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
								row_cells[0].width = Mm(width_list[0])

								row_cells[1].text = member.author.get_full_name()
								row_cells[1].width = Mm(width_list[1])

								if 'type_1' in request.POST or 'type_2' in request.POST:
									row_cells[2].text = member.name
									if member.technique:
										row_cells[2].text += ',\n' + member.technique
								else:
									row_cells[2].text = member.name_1
									if member.composer_1:
										row_cells[2].text += ',\n муз. ' + member.composer_1
									if member.poet_1:
										row_cells[2].text += ',\n сл. ' + member.poet_1

									row_cells[2].text += '\n\n' + member.name_2
									if member.composer_2:
										row_cells[2].text += ',\n муз. ' + member.composer_2
									if member.poet_2:
										row_cells[2].text += ',\n сл. ' + member.poet_2

								row_cells[2].width = Mm(width_list[2])

								row_cells[3].text =  ''
								if co_teachers_docx[member.pk]:
									teachers_flag = False
									for co_teacher in co_teachers_docx[member.pk]:
										if teachers_flag:
											row_cells[3].text += '\n'
										row_cells[3].text += co_teacher.short_profile_type() + ' ' + co_teacher.get_file_name()
										teachers_flag = True

								row_cells[3].width = Mm(width_list[3])

								row_cells[4].text = str(member.author.main_user.profile.get_institute_full())
								row_cells[4].width = Mm(width_list[4])

								if 'type_3' in request.POST:
									row_cells[5].text = member.descritpion
									row_cells[5].width = Mm(width_list[5])

							document.add_paragraph()

			response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
			response['Content-Disposition'] = 'inline; filename=' + file_name + ' (' + dte.strftime('%d-%m-%Y') + ').docx'
			document.save(response)

			return response


	args = {
		'near_pictures': near_pictures,
		'far_pictures': far_pictures,
		'near_moves': near_moves,
		'far_moves': far_moves,

		'nomination_picture_list': nomination_picture_list,
		'nomination_move_list': nomination_move_list,

		'near_pictures_nominations': near_pictures_nominations,
		'far_pictures_nominations': far_pictures_nominations,
		'near_moves_nominations': near_moves_nominations,
		'far_moves_nominations': far_moves_nominations,

		'co_teachers': co_teachers,
		'co_teachers_vocal': co_teachers_vocal
	}

	return render(request, 'statistic_contestant_add.html', args)


@login_required(login_url='/login/')
def statistic_contestant_view(request, part = None):
	if not request.user.profile.admin_access:
		return redirect('home')

	picture_list = {}
	move_list = {}
	invoice_list = {}
	protocol_list = {}
	statement_list = {}
	personal_list = {}
	zip_list = {}

	users = User.objects.filter(is_active = True, profile__member_access = True).exclude(username = 'admin')
	users_not_active = User.objects.filter(is_active = False, profile__member_access = True).exclude(username = 'admin')

	for user in users:
		picture_list[user.pk] = Picture.objects.filter(author__main_user = user).order_by('participation', 'nomination', 'author__team', 'author__surname', 'author__name', 'author__name2')
		move_list[user.pk] = Movie.objects.filter(author__main_user = user).order_by('participation', 'nomination', 'author__team', 'author__surname', 'author__name', 'author__name2')
		invoice_list[user.pk] = Invoice.objects.filter(payer = user)
		protocol_list[user.pk] = Protocol.objects.filter(owner = user)
		statement_list[user.pk] = Statement.objects.filter(owner = user)
		personal_list[user.pk] = Personal.objects.filter(owner = user)

		if invoice_list[user.pk] or protocol_list[user.pk] or statement_list[user.pk] or personal_list[user.pk]:
			zip_list[user.pk] = True
		else:
			zip_list[user.pk] = False

	if request.POST:
		if 'type_1' in request.POST:
			locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

			dte = date.today()
			document = Document()
			section = document.sections[-1]
			new_width, new_height = section.page_height, section.page_width
			section.orientation = WD_ORIENT.LANDSCAPE
			section.page_width = Mm(297)
			section.page_height = Mm(210)
			section.left_margin = Mm(30)
			section.right_margin = Mm(10)
			section.top_margin = Mm(10)
			section.bottom_margin = Mm(10)
			section.header_distance = Mm(10)
			section.footer_distance = Mm(10)

			style = document.styles['Normal']
			font = style.font
			font.name = 'Times New Roman'
			font.size = Pt(12)


			document.add_paragraph('Список учреждений').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			p = document.add_paragraph()
			p.add_run(dte.strftime('%d %B %Y')).italic = True
			p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

			table = document.add_table(rows=1, cols=5)
			table.allow_autifit = False
			table.style = 'Table Grid'
			table.columns[0].width = Mm(10)
			table.columns[1].width = Mm(130)
			table.columns[2].width = Mm(77)
			table.columns[3].width = Mm(20)
			table.columns[4].width = Mm(20)


			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = '№'
			hdr_cells[0].paragraphs[0].runs[0].font.bold = True
			hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[0].width = Mm(10)
			hdr_cells[1].text = 'Учреждение'
			hdr_cells[1].paragraphs[0].runs[0].font.bold = True
			hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[1].width = Mm(130)
			hdr_cells[2].text = 'Администратор/Email/Телефон'
			hdr_cells[2].paragraphs[0].runs[0].font.bold = True
			hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[2].width = Mm(77)
			hdr_cells[3].text = 'Кол-во работ вокал'
			hdr_cells[3].paragraphs[0].runs[0].font.bold = True
			hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[3].width = Mm(20)
			hdr_cells[4].text = 'Кол-во работ ДПИ'
			hdr_cells[4].paragraphs[0].runs[0].font.bold = True
			hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			hdr_cells[4].width = Mm(20)


			cnt = 0
			for user in users:
				cnt += 1
				row_cells = table.add_row().cells
				row_cells[0].text = str(cnt)
				row_cells[0].paragraphs[0].runs[0].font.bold = True
				row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				row_cells[0].width = Mm(10)

				row_cells[1].text = user.profile.get_institute_full() + '\n(' + user.profile.get_category_display() + ')'
				row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				row_cells[1].width = Mm(130)

				row_cells[2].text = user.profile.get_full_name() + '\n' + user.email + '\n' + user.profile.phone
				row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				row_cells[2].width = Mm(77)

				row_cells[3].text = str(move_list[user.pk].count())
				row_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				row_cells[3].width = Mm(20)

				row_cells[4].text = str(picture_list[user.pk].count())
				row_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
				row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
				row_cells[4].width = Mm(20)


			response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
			response['Content-Disposition'] = 'attachment; filename=Institutions (' + dte.strftime('%d-%m-%Y') + ').docx'
			document.save(response)

			return response

		if 'type_2' in request.POST:
			return redirect('profiles:view_new_admin_profile')


	args = {
		'users': users,
		'users_not_active': users_not_active,
		'picture_list': picture_list,
		'move_list': move_list,
		'invoice_list': invoice_list,
		'protocol_list': protocol_list,
		'statement_list': statement_list,
		'personal_list': personal_list,
		'zip_list': zip_list
	}
	return render(request, 'statistic_contestant.html', args)


@login_required(login_url='/login/')
def statistic_invoices_view(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	profiles = Profile.objects.filter(member_access=True).exclude(user__username='admin').exclude(user__is_active=False).order_by('surname', 'name', 'name2')
	invoices = Invoice.objects.all()
	array = {}
	array_sum = {}
	for profile in profiles:
		inces = invoices.filter(payer=profile.user)
		if inces:
			summa = 0
			array[profile.id] = inces
			for ince in inces:
				summa += ince.summa
			array_sum[profile.id] = summa


	if request.POST:
		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()
		document = Document()
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = Mm(297)
		section.page_height = Mm(210)
		section.left_margin = Mm(30)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)


		document.add_paragraph('Список оплат').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		p = document.add_paragraph()
		p.add_run(dte.strftime('%d.%b.%Y')).italic = True
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

		table = document.add_table(rows=1, cols=5)
		table.allow_autifit = False
		table.style = 'TableGrid'
		table.columns[0].width = Mm(10)
		table.columns[1].width = Mm(70)
		table.columns[2].width = Mm(70)
		table.columns[3].width = Mm(30)
		table.columns[4].width = Mm(77)


		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = '№'
		hdr_cells[0].width = Mm(10)
		hdr_cells[1].text = 'Конкурсант'
		hdr_cells[1].width = Mm(70)
		hdr_cells[2].text = 'Учреждение'
		hdr_cells[2].width = Mm(70)
		hdr_cells[3].text = 'Оплаченная сумма'
		hdr_cells[3].width = Mm(30)
		hdr_cells[4].text = 'Квитанции'
		hdr_cells[4].width = Mm(77)

		cnt = 1
		for profile in profiles:
			row_cells = table.add_row().cells
			row_cells[0].text = str(cnt)
			row_cells[0].width = Mm(10)
			row_cells[1].text = profile.get_full_name()
			row_cells[1].width = Mm(70)
			row_cells[2].text = profile.institution
			row_cells[2].width = Mm(70)
			row_cells[3].text = '0'
			if profile.id in array_sum:
				row_cells[3].text = str(array_sum[profile.id])
			row_cells[3].width = Mm(30)
			row_cells[4].text = ''
			if profile.id in array:
				for invoice in array[profile.id]:
					row_cells[4].text += 'Дата: ' + str(invoice.date) + ', сумма: ' + str(invoice.summa) + ' руб.\n'
			row_cells[4].width = Mm(77)

			cnt += 1



		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=InvoiceList (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response

	args = {
		'profiles': profiles,
		'array': array,
		'array_sum': array_sum
	}
	return render(request, 'statistic_invoices.html', args)


@login_required(login_url='/login/')
def juri_view(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	chef_juris = Profile.objects.filter(chef_juri_accecc=True) 
	juris = Profile.objects.filter(juri_accecc=True).exclude(chef_juri_accecc=True)


	if request.method=='POST':
		if 'new' in request.POST:
			return redirect('juri_new')

		return redirect('juri_set')

	args = {
		'chef_juris': chef_juris,
		'juris': juris,
	}
	return render(request, 'view_juri.html', args)


@login_required(login_url='/login/')
def juri_set(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	pretendents = Profile.objects.filter(user__is_active=True).exclude(user__username='admin').exclude(juri_accecc=True).exclude(chef_juri_accecc=True).order_by('surname', 'name', 'name2')

	if request.POST:
		form = SetJuriForm(request.POST, pretendents=pretendents)	
		if form.is_valid():
			pretendent = form.cleaned_data['juri']
			pretendent.juri_accecc = True
			pretendent.chef_juri_accecc = form.cleaned_data['chef']
			pretendent.juri_type = form.cleaned_data['juri_type']
			pretendent.save()

			return redirect('juri_view')

	form = SetJuriForm(pretendents=pretendents)
	args = {
		'form': form
	}
	return render(request, 'juri_set.html', args)


@login_required(login_url='/login/')
def juri_change(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	pretendent = Profile.objects.get(pk = pk)

	if request.POST:
		form = ChangeJuriForm(request.POST, instance=pretendent)	
		if form.is_valid():
			pretendent = form.save(commit=False)
			pretendent.save()

			return redirect('juri_view')

	form = ChangeJuriForm(instance=pretendent)
	args = {
		'pretendent': pretendent,
		'form': form
	}
	return render(request, 'juri_change.html', args)


@login_required(login_url='/login/')
def juri_new(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	if request.method=='POST':
		form = NewJuriForm(request.POST)
		if form.is_valid():
			password = '1QaZ2WsX3EdC'#User.objects.make_random_password(8)
			new_user = form.save(commit=False)
			new_user.username = new_user.email
			new_user.is_active = True
			new_user.set_password(password)
			new_user.save()

			new_user.profile.name = form.cleaned_data['name']
			new_user.profile.name2 = form.cleaned_data['name2']
			new_user.profile.surname = form.cleaned_data['surname']
			new_user.profile.institution = form.cleaned_data['institution']
			new_user.profile.adress = form.cleaned_data['adress']
			new_user.profile.rank = form.cleaned_data['rank']
			new_user.profile.member_access = False
			new_user.profile.juri_accecc = True
			new_user.profile.juri_type = form.cleaned_data['juri_type']
			new_user.profile.chef_juri_accecc = form.cleaned_data['chef_juri_accecc']
			new_user.profile.save()

			protocol = 'http'
			if request.is_secure():
				protocol = 'https'
			current_site = get_current_site(request)

			mail_subject = 'Конкурс "Гавриловские гуляния"'
			to_email = new_user.email
			sex = new_user.profile.sex()
			sex_valid = new_user.profile.sex_valid()
			signature = 'С уважением,\r\nАдминистрация конкурса'
			sign = signature.split('\r\n')

			args = {
				'sex': sex,
				'sex_valid': sex_valid,
				'name': new_user.profile.get_io_name(),
				'protocol': protocol,
				'domain': current_site.domain,
				'signature': signature,
				'sign': sign,
				'email': to_email,
				'password': password,
				'chef': new_user.profile.chef_juri_accecc
			}
				
			message = render_to_string('juri_active_email.html', args)

			message_html = render_to_string('juri_active_email_html.html', args)

			send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

			return redirect('juri_view')

		args = {
			'form': form
		}
		return render(request, 'juri_new.html', args)

	form = NewJuriForm()
	args = {
		'form': form
	}
	return render(request, 'juri_new.html', args)




@login_required(login_url='/login/')
def view_contestant(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	contestant = Profile.objects.get(pk = pk)
	pictures = Picture.objects.filter(author = contestant.user)
	movies = Movie.objects.filter(author = contestant.user)

	copicture_list = {}
	comovie_list = {}

	for picture in pictures:
		copicture_list[picture.pk] = CoPicturee.objects.filter(picture = picture, coauthor__main_user = contestant.user)

	for movie in movies:
		comovie_list[movie.pk] = CoMovie.objects.filter(movie = movie, coauthor__main_user = contestant.user)

	coprofiles = CoProfile.objects.filter(main_user = contestant.user).order_by('main_user')
	
	args = {
		'contestant': contestant,
		'pictures': pictures,
		'movies': movies,
		'copicture_list': copicture_list,
		'comovie_list': comovie_list,
		'coprofiles': coprofiles
	}
	return render(request, 'view_contestant.html', args)


@login_required(login_url='/login/')
def move_order_view(request):
	if not request.user.profile.admin_access:
		return redirect('home')


	user_list = list(Movie.objects.filter(participation = '1').order_by('nomination').distinct().values_list('author__main_user', flat=True))
	users = User.objects.filter(pk__in = user_list)

	movie_list = {}
	for user in users:
		movie_list[user.pk] = Movie.objects.filter(participation = '1', author__main_user = user)

	movies = Movie.objects.filter(participation = '1').order_by('nomination', 'author__main_user__profile__category')

	co_movies = {}
	for movie in movies:
		co_movies[movie.pk] = CoMovie.objects.filter(movie = movie)

	scene_numbers = movies.filter(scene_num=None)
	scene_numbers_old = movies.filter(scene_num__isnull=False).order_by('-scene_num')

	cnt = 1
	if scene_numbers_old:
		cnt = int(scene_numbers_old[0].scene_num)
		cnt += 1

	if scene_numbers:
		for movie in scene_numbers:
			movie.scene_num = cnt
			movie.save()
			cnt += 1

	movies = Movie.objects.filter(participation = '1').order_by('scene_num')


	if request.POST:
		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()
		file_name_add = ''

		document = Document()
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = Mm(297)
		section.page_height = Mm(210)
		section.left_margin = Mm(30)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)


		document.add_paragraph('Порядок выступлений').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		p = document.add_paragraph()
		p.add_run(dte.strftime('%d %B %Y')).italic = True
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT


		table = document.add_table(rows=1, cols=5)
		table.allow_autifit = False
		table.style = 'TableGrid'
		table.columns[0].width = Mm(10)
		table.columns[1].width = Mm(70)
		table.columns[2].width = Mm(70)
		table.columns[3].width = Mm(60)
		table.columns[4].width = Mm(47)


		#257

		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = '№'
		hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[0].width = Mm(10)
		hdr_cells[1].text = 'Название'
		hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[1].width = Mm(70)
		hdr_cells[2].text = 'Конкурсант\n(коллектив)'
		hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[2].width = Mm(70)
		hdr_cells[3].text = 'Номинация'
		hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[3].width = Mm(60)
		hdr_cells[4].text = 'Учреждение'
		hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[4].width = Mm(47)


		cnt = 1
		for movie in movies:
			row_cells = table.add_row().cells
			row_cells[0].text = str(cnt)
			row_cells[0].paragraphs[0].runs[0].font.bold = True
			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(10)

			row_cells[1].text = movie.name_1
			if movie.composer_1:
				row_cells[1].text += ' муз. ' +  movie.composer_1
			if  movie.poet_1:
				row_cells[1].text += ' сл. ' +  movie.poet_1

			row_cells[1].text += '\n' + movie.name_2
			if movie.composer_2:
				row_cells[1].text += ' муз. ' +  movie.composer_2
			if  movie.poet_2:
				row_cells[1].text += ' сл. ' +  movie.poet_2
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(70)

			row_cells[2].text = movie.author.get_full_name()
			if co_movies[movie.pk]:
				for co_movie in co_movies[movie.pk]:
					row_cells[2].text += '\n' + co_movie.coauthor.short_profile_type() + ' ' + co_movie.coauthor.get_file_name()
			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[2].width = Mm(70)

			row_cells[3].text = str(movie.nomination)
			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[3].width = Mm(60)

			row_cells[4].text = movie.author.main_user.profile.get_institute() + '\n' + movie.author.main_user.profile.short_category_type()
			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[4].width = Mm(47)

			cnt += 1

		file_name = 'OrderList'
		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=' + file_name +' (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response



	args = {
		'movies': movies,
		'co_movies': co_movies
	}
	return render(request, 'move_order.html', args)



@login_required(login_url='/login/')
def get_contestant_zip_view(request, pk=None):
	if not request.user.profile.admin_access:
		return redirect('home')

	user = User.objects.filter(pk = pk).first()
	if not user:
		return HttpResponse('Пользователь не найден')

	invoice_list = Invoice.objects.filter(payer = user)
	protocol_list = Protocol.objects.filter(owner = user)
	statement_list = Statement.objects.filter(owner = user)
	personal_list = Personal.objects.filter(owner = user)


	f = io.BytesIO()
	zip_title = translit.slugify(user.profile.get_institute_zip())+'.zip'
	zip_arch = ZipFile( f, 'a' )

	cnt = 1
	for invoice in invoice_list:
		docx_title = 'Квитанция ' + str(cnt)
		filename, fileext = os.path.splitext(invoice.file.name)
		doc = settings.MEDIA_ROOT + '/' + invoice.file.name
		zip_arch.write(doc, docx_title + fileext)

		cnt += 1

	cnt = 1
	for protocol in protocol_list:
		docx_title = 'Протокол ' + str(cnt)
		filename, fileext = os.path.splitext(protocol.file.name)
		doc = settings.MEDIA_ROOT + '/' + protocol.file.name
		zip_arch.write(doc, docx_title + fileext)

		cnt += 1

	cnt = 1
	for statement in statement_list:
		docx_title = 'Заявка ' + str(cnt)
		filename, fileext = os.path.splitext(statement.file.name)
		doc = settings.MEDIA_ROOT + '/' + statement.file.name
		zip_arch.write(doc, docx_title + fileext)

		cnt += 1

	cnt = 1
	for personal in personal_list:
		docx_title = 'Согласие ' + str(cnt)
		filename, fileext = os.path.splitext(personal.file.name)
		doc = settings.MEDIA_ROOT + '/' + personal.file.name
		zip_arch.write(doc, docx_title + fileext)

		cnt += 1

	zip_arch.close()
	response = HttpResponse(f.getvalue(), content_type='application/zip')
	response['Content-Disposition'] = 'attachment; filename=' + zip_title
	return response


# --------------------------------
#           Для ajax'а
# --------------------------------
@login_required(login_url='/login/')
def delete_juri(request):
	juri_pk = request.GET['juri']
	juri = Profile.objects.get(pk=juri_pk)

	juri.juri_accecc = False
	juri.chef_juri_accecc = False
	juri.save()

	return HttpResponse(True)


def ajax_activate(request):
	pk = int(request.GET['pk'])
	user = User.objects.get(pk = pk)
	user.is_active = True
	user.save()


	picture_list = {}
	move_list = {}

	users = User.objects.filter(is_active = True, profile__member_access = True).exclude(username = 'admin')
	users_not_active = User.objects.filter(is_active = False, profile__member_access = True).exclude(username = 'admin')

	for user in users:
		picture_list[user.pk] = Picture.objects.filter(author__main_user = user).order_by('participation', 'nomination', 'author__team', 'author__surname', 'author__name', 'author__name2')
		move_list[user.pk] = Movie.objects.filter(author__main_user = user).order_by('participation', 'nomination', 'author__team', 'author__surname', 'author__name', 'author__name2')

	args = {
		'users': users,
		'users_not_active': users_not_active,
		'picture_list': picture_list,
		'move_list': move_list
	}

	data = render_to_string('view_profiles_refresh.html', args)
	return HttpResponse(data)
