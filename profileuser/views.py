import os
import datetime
import locale

from datetime import date

from django.conf import settings
from django.core.files.storage import default_storage

from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.contrib.auth.decorators import login_required
from django.contrib.auth import update_session_auth_hash

from django.contrib.sites.shortcuts import get_current_site

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Pt

from django.contrib.auth.models import User

from django.template.loader import render_to_string
from django.core.mail import EmailMessage, send_mail

from .forms import ProfileUdpateForm, CoProfileForm, CoProfileTeamForm, CoProfileTeamEditForm, UserRegistrationForm

from .models import Profile, CoProfile
from movies.models import Movie


@login_required(login_url='/login/')
def view_edit_profile(request):
	username = request.user.username

	if request.method=='POST':
		form_profile = ProfileUdpateForm(request.POST, instance=request.user.profile, label_suffix=':')

		if "passchange" in request.POST:
			return redirect('passchange')

		if form_profile.is_valid():
			if form_profile.has_changed():
				profile_form = form_profile.save(False)
				profile_form.save()	

			
			args ={
				'form': form_profile,
				'done_flag': True
			}
			return render(request, 'profileuser/view_edit_profile.html', args)

		args ={
			'form': form_profile,
			'done_flag': False
		}
		return render(request, 'profileuser/view_edit_profile.html', args)

	form_profile = ProfileUdpateForm(instance=request.user.profile, label_suffix=':')

	args = {
		'form': form_profile, 
		'done_flag': False
	}
	return render(request, 'profileuser/view_edit_profile.html', args)


@login_required(login_url='/login/')
def view_new_admin_profile(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	if request.method=='POST':
		user_form = UserRegistrationForm(request.POST)
		if user_form.is_valid():
			password = User.objects.make_random_password(8)
			new_user = user_form.save(commit=False)
			new_user.username = new_user.email
			new_user.is_active = True
			new_user.set_password(password)
			new_user.save()

			new_user.profile.less_institution = user_form.cleaned_data['less_institution']
			new_user.profile.category = user_form.cleaned_data['category']
			new_user.profile.surname = user_form.cleaned_data['surname']
			new_user.profile.name = user_form.cleaned_data['name']
			new_user.profile.name2 = user_form.cleaned_data['name2']
			new_user.profile.profile_type = user_form.cleaned_data['profile_type']
			new_user.profile.phone = user_form.cleaned_data['phone']
			new_user.profile.institution = user_form.cleaned_data['institution']
			new_user.profile.institution_shot = user_form.cleaned_data['institution_shot']
			new_user.profile.adress = user_form.cleaned_data['adress']

			new_user.profile.save()

			if user_form.cleaned_data['add_team']:
				CoProfile.objects.create(
					main_user = new_user,
					coprofile_type = '2',
					team = user_form.cleaned_data['team']
				)

			protocol = 'http'
			if request.is_secure():
				protocol = 'https'
			current_site = get_current_site(request)

			mail_subject = 'Конкурс "Гавриловские гуляния"'
			to_email = new_user.email
			sex = new_user.profile.sex()
			sex_valid = new_user.profile.sex_valid()
			signature = 'С уважением,\r\nАдминистрация конкурса'
			sign = signature.split('\r\n')

			args = {
				'sex': sex,
				'sex_valid': sex_valid,
				'name': new_user.profile.get_io_name(),
				'protocol': protocol,
				'domain': current_site.domain,
				'signature': signature,
				'sign': sign,
				'email': to_email,
				'password': password,
			}

			message = render_to_string('profileuser/new_profile_email.html', args)

			message_html = render_to_string('profileuser/new_profile_email_html.html', args)

			send_mail(mail_subject, message, settings.EMAIL_HOST_USER, [to_email], fail_silently=True, html_message=message_html)

			return redirect('statistic_contestant')

		return render(request, 'profileuser/register.html', {'form': user_form})

	user_form = UserRegistrationForm()
	return render(request, 'profileuser/register.html', {'form': user_form})


@login_required(login_url='/login/')
def view_edit_profile_admin(request, pk = None):
	user = User.objects.get(pk = pk)
	username = user.profile.get_institute_zip()

	coprofiles = CoProfile.objects.filter(main_user = user, self_flag = False, profile_type__in = ['1', '2', '3', '4', '5']).order_by('surname', 'name', 'name2')
	teams = CoProfile.objects.filter(main_user = user, self_flag = False, profile_type = '0').order_by('team', 'surname', 'name', 'name2')

	if request.method=='POST':
		if 'add-teacher' in request.POST:
			return redirect('profiles:new_coprofile_admin', pk = pk)
		if 'add-team' in request.POST:
			return redirect('profiles:new_team_coprofile_admin', pk = pk)
		form_profile = ProfileUdpateForm(request.POST, instance=user.profile, label_suffix=':')


		if form_profile.is_valid():
			if form_profile.has_changed():
				profile_form = form_profile.save(False)
				profile_form.save()


			args ={
				'form': form_profile,
				'done_flag': True,
				'coprofiles': coprofiles,
				'teams': teams,
				'username': username
			}
			return render(request, 'profileuser/view_edit_profile_admin.html', args)

		args ={
			'form': form_profile,
			'done_flag': False,
			'coprofiles': coprofiles,
			'teams': teams,
			'username': username
		}
		return render(request, 'profileuser/view_edit_profile_admin.html', args)

	form_profile = ProfileUdpateForm(instance=user.profile, label_suffix=':')

	args = {
		'form': form_profile,
		'done_flag': False,
		'coprofiles': coprofiles,
		'teams': teams,
		'username': username
	}
	return render(request, 'profileuser/view_edit_profile_admin.html', args)


@login_required(login_url='/login/')
def view_coprofiles(request):
	if not request.user.profile.member_access:
		return redirect('home')
		
	if request.method=='POST':
		return redirect('profiles:new_coprofile')

	coprofiles = CoProfile.objects.filter(main_user = request.user, self_flag = False, profile_type__in = ['1', '2', '3', '4', '5']).order_by('main_user')
	
	args = {
		'coprofiles': coprofiles,
	}
	return render(request, 'profileuser/view_coprofiles.html', args)


@login_required(login_url='/login/')
def new_coprofile(request):
	if not request.user.profile.member_access:
		return redirect('home')

	if request.method=='POST':
		form_coprofile = CoProfileForm(request.POST, label_suffix='')

		if form_coprofile.is_valid():
			coprofile = form_coprofile.save(commit=False)
			coprofile.main_user = request.user
			coprofile.save()	

			return redirect('profiles:view_coprofiles')

		args ={
			'form': form_coprofile, 
		}
		return render(request, 'profileuser/new_coprofile.html', args)

	form_coprofile = CoProfileForm(label_suffix='')

	args = {
		'form': form_coprofile, 
	}
	return render(request, 'profileuser/new_coprofile.html', args)


@login_required(login_url='/login/')
def new_coprofile_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	main_user = User.objects.get(pk = pk)
	username = main_user.profile.get_institute_zip()

	if request.method=='POST':
		form_coprofile = CoProfileForm(request.POST, label_suffix='')

		if form_coprofile.is_valid():
			coprofile = form_coprofile.save(commit=False)
			coprofile.main_user = main_user
			coprofile.save()	

			return redirect('profiles:view_edit_profile_admin', pk = pk)

		args ={
			'form': form_coprofile, 
			'username': username
		}
		return render(request, 'profileuser/new_coprofile.html', args)

	form_coprofile = CoProfileForm(label_suffix='')

	args = {
		'form': form_coprofile, 
		'username': username
	}
	return render(request, 'profileuser/new_coprofile.html', args)


@login_required(login_url='/login/')
def view_edit_coprofile(request, pk):
	if not request.user.profile.member_access:
		return redirect('home')

	coprofile = CoProfile.objects.get(pk = pk)

	if request.method=='POST':
		form_coprofile = CoProfileForm(request.POST, instance=coprofile, label_suffix=':')

		if form_coprofile.is_valid():
			if form_coprofile.has_changed():
				coprofile = form_coprofile.save(False)
				coprofile.save()	

			return redirect('profiles:view_coprofiles')

		args ={
			'form': form_coprofile, 
		}
		return render(request, 'profileuser/view_edit_coprofile.html', args)

	form_coprofile = CoProfileForm(instance=coprofile, label_suffix=':')

	args = {
		'form': form_coprofile, 
	}
	return render(request, 'profileuser/view_edit_coprofile.html', args)


@login_required(login_url='/login/')
def view_edit_coprofile_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	coprofile = CoProfile.objects.get(pk = pk)
	username = coprofile.main_user.profile.get_institute_zip()

	if request.method=='POST':
		form_coprofile = CoProfileForm(request.POST, instance=coprofile, label_suffix=':')

		if form_coprofile.is_valid():
			if form_coprofile.has_changed():
				coprofile = form_coprofile.save(False)
				coprofile.save()	

			return redirect('profiles:view_edit_profile_admin', pk = coprofile.main_user.profile.pk)

		args ={
			'form': form_coprofile, 
			'username': username
		}
		return render(request, 'profileuser/view_edit_coprofile.html', args)

	form_coprofile = CoProfileForm(instance=coprofile, label_suffix=':')

	args = {
		'form': form_coprofile, 
		'username': username
	}
	return render(request, 'profileuser/view_edit_coprofile.html', args)


@login_required(login_url='/login/')
def view_team_coprofiles(request):
	if not request.user.profile.member_access:
		return redirect('home')

	if request.method=='POST':
		return redirect('profiles:new_team_coprofile')

	coprofiles = CoProfile.objects.filter(main_user = request.user, self_flag = False, profile_type = '0').order_by('main_user')

	args = {
		'coprofiles': coprofiles,
	}
	return render(request, 'profileuser/view_team_coprofiles.html', args)


@login_required(login_url='/login/')
def new_team_coprofile(request):
	if not request.user.profile.member_access:
		return redirect('home')

	if request.method=='POST':
		form_coprofile = CoProfileTeamForm(request.POST, label_suffix='')

		if form_coprofile.is_valid():
			coprofile = form_coprofile.save(commit=False)
			coprofile.main_user = request.user
			coprofile.save()

			return redirect('profiles:view_team_coprofiles')

		args ={
			'form': form_coprofile,
		}
		return render(request, 'profileuser/new_team_coprofile.html', args)

	form_coprofile = CoProfileTeamForm(label_suffix='')

	args = {
		'form': form_coprofile,
	}
	return render(request, 'profileuser/new_team_coprofile.html', args)


@login_required(login_url='/login/')
def new_team_coprofile_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	main_user = User.objects.get(pk = pk)
	username = main_user.profile.get_institute_zip()

	if request.method=='POST':
		form_coprofile = CoProfileTeamForm(request.POST, label_suffix='')

		if form_coprofile.is_valid():
			coprofile = form_coprofile.save(commit=False)
			coprofile.main_user = main_user
			coprofile.save()

			return redirect('profiles:view_edit_profile_admin', pk = pk)

		args ={
			'form': form_coprofile,
			'username': username
		}
		return render(request, 'profileuser/new_team_coprofile.html', args)

	form_coprofile = CoProfileTeamForm(label_suffix='')

	args = {
		'form': form_coprofile,
		'username': username
	}
	return render(request, 'profileuser/new_team_coprofile.html', args)


@login_required(login_url='/login/')
def view_edit_team_coprofile(request, pk):
	if not request.user.profile.member_access:
		return redirect('home')

	coprofile = CoProfile.objects.get(pk = pk)

	if request.method=='POST':
		form_coprofile = CoProfileTeamEditForm(request.POST, instance=coprofile, label_suffix=':')

		if form_coprofile.is_valid():
			if form_coprofile.has_changed():
				coprofile = form_coprofile.save(False)
				coprofile.save()

			return redirect('profiles:view_team_coprofiles')

		args ={
			'form': form_coprofile,
		}
		return render(request, 'profileuser/view_edit_team_coprofile.html', args)

	form_coprofile = CoProfileTeamEditForm(instance=coprofile, label_suffix=':')

	args = {
		'form': form_coprofile,
	}
	return render(request, 'profileuser/view_edit_team_coprofile.html', args)


@login_required(login_url='/login/')
def view_edit_team_coprofile_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	coprofile = CoProfile.objects.get(pk = pk)
	username = coprofile.main_user.profile.get_institute_zip()

	if request.method=='POST':
		form_coprofile = CoProfileTeamEditForm(request.POST, admin_access = True, instance=coprofile, label_suffix=':')

		if form_coprofile.is_valid():
			if form_coprofile.has_changed():
				coprofile = form_coprofile.save(False)
				coprofile.save()

			return redirect('profiles:view_edit_profile_admin', pk = coprofile.main_user.profile.pk)

		args ={
			'form': form_coprofile,
			'username': username
		}
		return render(request, 'profileuser/view_edit_team_coprofile.html', args)

	form_coprofile = CoProfileTeamEditForm(admin_access = True, instance=coprofile, label_suffix=':')

	args = {
		'form': form_coprofile,
		'username': username
	}
	return render(request, 'profileuser/view_edit_team_coprofile.html', args)


@login_required(login_url='/login/')
def view_comings(request):
	if not request.user.profile.register_accecc:
		return redirect('home')

	user_list = list(Movie.objects.filter(participation = '1').order_by('nomination').distinct().values_list('author__main_user', flat=True))
	users = User.objects.filter(pk__in = user_list)

	movie_list = {}
	for user in users:
		movie_list[user.pk] = Movie.objects.filter(participation = '1', author__main_user = user)

	movies = Movie.objects.filter(participation = '1').order_by('nomination', 'author__main_user__profile__category')
	
	scene_numbers = movies.filter(scene_num=None)
	scene_numbers_old = movies.filter(scene_num__isnull=False).order_by('-scene_num')

	cnt = 1
	if scene_numbers_old:
		cnt = int(scene_numbers_old[0].scene_num)
		cnt += 1

	if scene_numbers:
		for movie in scene_numbers:
			movie.scene_num = cnt
			movie.save()
			cnt += 1

	movies = Movie.objects.filter(participation = '1', has_come = True).order_by('scene_num')

	if request.POST:
		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()
		file_name_add = ''

		document = Document()
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = Mm(297)
		section.page_height = Mm(210)
		section.left_margin = Mm(30)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)


		document.add_paragraph('Порядок выступлений').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		p = document.add_paragraph()
		p.add_run(dte.strftime('%d %B %Y')).italic = True
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT


		table = document.add_table(rows=1, cols=5)
		table.allow_autifit = False
		table.style = 'TableGrid'
		table.columns[0].width = Mm(10)
		table.columns[1].width = Mm(70)
		table.columns[2].width = Mm(70)
		table.columns[3].width = Mm(60)
		table.columns[4].width = Mm(47)


		#257

		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = '№'
		hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[0].width = Mm(10)
		hdr_cells[1].text = 'Название'
		hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[1].width = Mm(70)
		hdr_cells[2].text = 'Конкурсант\n(коллектив)'
		hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[2].width = Mm(70)
		hdr_cells[3].text = 'Номинация'
		hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[3].width = Mm(60)
		hdr_cells[4].text = 'Учреждение'
		hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[4].width = Mm(47)


		cnt = 1
		for movie in movies:
			row_cells = table.add_row().cells
			row_cells[0].text = str(cnt)
			row_cells[0].paragraphs[0].runs[0].font.bold = True
			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(10)
			row_cells[1].text = str(movie)
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(70)
			row_cells[2].text = movie.author.get_full_name()
			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[2].width = Mm(70)
			row_cells[3].text = str(movie.nomination)
			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[3].width = Mm(60)
			row_cells[4].text = movie.author.main_user.profile.get_institute_zip() + '\n'
			row_cells[4].text += '(' + movie.author.main_user.profile.get_category_display() + ')'
			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[4].width = Mm(47)

			cnt += 1

		file_name = 'OrderList (current)'
		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=' + file_name +' (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response



	args ={
		'users': users,
		'movie_list': movie_list
	}
	return render(request, 'profileuser/view_comings.html', args)


# --------------------------------
#           Для ajax'а
# --------------------------------
@login_required(login_url='/login/')
def ajax_del_coprofile(request):
	coprofile_pk = request.GET['coprofile']
	coprofile = CoProfile.objects.get(pk=coprofile_pk)

	if not request.user.profile.admin_access and coprofile.main_user != request.user:
		return HttpResponse(False)

	coprofile.delete()

	return HttpResponse(True)


def change_come_flag(request):
	pk = request.GET['pk']
	checked = request.GET['checked']

	movie = Movie.objects.get(pk = pk)
	if checked == '1':
		movie.has_come = True
	else:
		movie.has_come = False
	movie.save()

	return HttpResponse(True)
