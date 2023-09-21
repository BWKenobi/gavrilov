import os
import datetime
import locale

from datetime import date

from django.conf import settings
from django.core.files.storage import default_storage

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Pt

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import update_session_auth_hash
from django.contrib.auth.models import User

from .forms import PictureUploadForm, PictureEditForm
from profileuser.models import Profile, CoProfile
from .models import Picture, CoPicturee
from profileuser.models import CoProfile

@login_required(login_url='/login/')
def view_arts(request):
	pictures = Picture.objects.filter(author__main_user=request.user)
	copictures_list = {}

	if request.method=='POST':
		if 'addpict' in request.POST:
			return redirect('pictures:load_image')

		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()

		document = Document()

		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.PORTRAIT
		section.page_width = Mm(210)
		section.page_height = Mm(297)
		section.left_margin = Mm(20)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(14)


		p = document.add_paragraph()
		p.add_run('ЗАЯВКА').bold = True
		p.paragraph_format.space_after = 0
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

		p = document.add_paragraph()
		p.add_run('на участие во Всероссийском фестивале-конкурсе').bold = True
		p.paragraph_format.space_after = 0
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

		p = document.add_paragraph()
		p.add_run('народного творчества «Гавриловские гуляния»').bold = True
		p.paragraph_format.space_after = 0
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

		p = document.add_paragraph()
		p.add_run('(выставочная программа)').bold = True
		p.paragraph_format.space_after = 0
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

		p = document.add_paragraph()
		p.add_run(dte.strftime('%d %B %Y') + ' года').italic = True
		p.paragraph_format.space_after = 0
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

		document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('Директору государственного автономного\nпрофессионального образовательного учреждения\n\
		# 	Московской области\n«Московский Губернский колледж искусств»\nХусеинову Р.А.').italic = False
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

		# document.add_paragraph().paragraph_format.space_after = 0

		# table = document.add_table(rows=1, cols=2)
		# table.allow_autifit = False
		# table.style = 'TableGrid'
		# table.columns[0].width = Mm(120)
		# table.columns[1].width = Mm(60)

		# row_cells = table.rows[0].cells
		# row_cells[0].text = 'Фамилия'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.surname
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Имя'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.name
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Отчество'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.name2
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Пол'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.get_sex_type_display()
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells = row_cells[0].merge(row_cells[1])

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Образование (9 классов, 11 классов, НПО, СПО, ВПО)'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.get_education_display()
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Название образовательного учреждения'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.education_name
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Место нахождения образовательного учреждения'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.education_address
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Документ об образовании (аттестат, диплом)'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.get_diploma_type_display()
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Серия и номер'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.series_number_diploma
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Дата выдачи'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = str(statement.student.profile.date_diploma)
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells = row_cells[0].merge(row_cells[1])

		# if statement.student.profile.art_school:
		# 	row_cells = table.add_row().cells
		# 	row_cells[0].text = 'Название образовательного учреждения (ДШИ/ДМШ)'
		# 	row_cells[0].width = Mm(120)
		# 	row_cells[1].text = statement.student.profile.art_name
		# 	row_cells[1].width = Mm(60)

		# 	row_cells = table.add_row().cells
		# 	row_cells[0].text = 'Место нахождения образовательного учреждения (ДШИ/ДМШ)'
		# 	row_cells[0].width = Mm(120)
		# 	row_cells[1].text = statement.student.profile.art_address
		# 	row_cells[1].width = Mm(60)

		# 	row_cells = table.add_row().cells
		# 	row_cells[0].text = 'Серия и номер свидетельства'
		# 	row_cells[0].width = Mm(120)
		# 	row_cells[1].text = statement.student.profile.series_number_art
		# 	row_cells[1].width = Mm(60)

		# 	row_cells = table.add_row().cells
		# 	row_cells[0].text = 'Дата выдачи свидетельства'
		# 	row_cells[0].width = Mm(120)
		# 	row_cells[1].text = str(statement.student.profile.date_art)
		# 	row_cells[1].width = Mm(60)

		# 	row_cells = table.add_row().cells
		# 	row_cells = row_cells[0].merge(row_cells[1])

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Дата рождения'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = str(statement.student.profile.birthday)
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Место рождения'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.birth_place
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Контактный телефон (личный)'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.phone
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Адрес электронной почты'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.email
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Гражданство'
		# row_cells[0].width = Mm(120)
		# if statement.student.profile.foreigner == '1':
		# 	row_cells[1].text = 'РФ'
		# else:
		# 	row_cells[1].text = statement.student.profile.nationality
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells = row_cells[0].merge(row_cells[1])

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Документ, удостоверяющий личность (название)'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = 'Паспорт'
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Серия, №'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.series_number
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Дата выдачи'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = str(statement.student.profile.date_pas)
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Кем выдан'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.department
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Код подразделения'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.number
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Адрес регистрации'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.address
		# row_cells[1].width = Mm(60)

		# row_cells = table.add_row().cells
		# row_cells = row_cells[0].merge(row_cells[1])

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'СНИЛС'
		# row_cells[0].width = Mm(120)
		# row_cells[1].text = statement.student.profile.snils
		# row_cells[1].width = Mm(60)

		# parents = Parent.objects.filter(student = statement.student)


		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('ЗАЯВЛЕНИЕ').bold = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

		# document.add_paragraph()

		# p = document.add_paragraph()
		# p.paragraph_format.first_line_indent = Mm(12.5)
		# p.add_run('Прошу принять документы и допустить меня к участию в конкурсе (и к вступительным испытаниям при их наличии):').bold = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


		# instruments = ''
		# if add_field != '':
		# 	instruments = ' (инструмент: ' + add_field + ')'


		# p = document.add_paragraph()
		# p.add_run('на специальность ').bold = True
		# p.add_run(speciality).underline = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('вид подготовки ').bold = True
		# p.add_run(type + instruments).underline = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('форма обучения ').bold = True
		# p.add_run(education_from).underline = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('- на места, финансируемые из бюджета Московской области ')
		# if budget:
		# 	p.add_run('Да').underline = True
		# else:
		# 	p.add_run('Нет').underline = True

		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('- по договору об оказании платных образовательных услуг ')
		# if budget:
		# 	p.add_run('Нет').underline = True
		# else:
		# 	p.add_run('Да').underline = True

		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.paragraph_format.first_line_indent = Mm(12.5)
		# p.add_run('Согласен(а) на обработку своих персональных данных в порядке, установленном Федеральным законом от 27.07.2006 № 152-ФЗ "О персональных данных".').bold = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('____________________________/______________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

		# p = document.add_paragraph()
		# p.add_run('	                    (подпись поступающего)                       (расшифровка)').italic = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# document.add_paragraph().paragraph_format.space_after = 0
		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('О себе сообщаю следующее: ').bold = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('Сведения об образовании:')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# table = document.add_table(rows=1, cols=2)
		# table.allow_autifit = False
		# table.style = 'TableGrid'
		# table.columns[0].width = Mm(160)
		# table.columns[1].width = Mm(20)

		# hdr_cells = table.rows[0].cells
		# hdr_cells[0].text = 'Название учреждения '
		# for paragraph in hdr_cells[0].paragraphs:
		# 	for run in paragraph.runs:
		# 		run.font.bold = True
		# hdr_cells[0].width = Mm(160)
		# hdr_cells[1].text = 'Да/Нет'
		# for paragraph in hdr_cells[1].paragraphs:
		# 	for run in paragraph.runs:
		# 		run.font.italic = True
		# 		run.font.bold = True
		# hdr_cells[1].width = Mm(20)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'общеобразовательное учреждение 9 классов'
		# row_cells[0].width = Mm(160)
		# if statement.student.profile.education == '1':
		# 	row_cells[1].text = 'Да'
		# 	for paragraph in row_cells[1].paragraphs:
		# 		paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# row_cells[1].width = Mm(20)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'общеобразовательное учреждение 11 классов'
		# row_cells[0].width = Mm(160)
		# if statement.student.profile.education == '2':
		# 	row_cells[1].text = 'Да'
		# 	for paragraph in row_cells[1].paragraphs:
		# 		paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# row_cells[1].width = Mm(20)

		# row_cells = table.add_row().cells
		# row_cells[0].text = ' образовательное учреждение начального профессионального образования (НПО) '
		# row_cells[0].width = Mm(160)
		# if statement.student.profile.education == '5':
		# 	row_cells[1].text = 'Да'
		# 	for paragraph in row_cells[1].paragraphs:
		# 		paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# row_cells[1].width = Mm(20)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'профессиональное образовательное учреждение (СПО) '
		# row_cells[0].width = Mm(160)
		# if statement.student.profile.education == '3':
		# 	row_cells[1].text = 'Да'
		# 	for paragraph in row_cells[1].paragraphs:
		# 		paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# row_cells[1].width = Mm(20)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'образовательное учреждение высшего образования'
		# row_cells[0].width = Mm(160)
		# if statement.student.profile.education == '4':
		# 	row_cells[1].text = 'Да'
		# 	for paragraph in row_cells[1].paragraphs:
		# 		paragraph.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		# row_cells[1].width = Mm(20)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Другое ______________________________ '
		# row_cells[0].width = Mm(160)
		# row_cells[1].width = Mm(20)

		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# if statement.student.profile.diploma_type == '1':
		# 	p.add_run('Аттестат (серия/номер): ')
		# else:
		# 	p.add_run('Диплом (серия/номер): ')
		# p.add_run(statement.student.profile.series_number_diploma).underline = True
		# p.add_run(', выдан ')
		# p.add_run(str(statement.student.profile.date_diploma) + ' г.').underline = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('Трудовой стаж (если есть) ______________ лет ________________ мес.')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('Иностранный язык ')
		# if statement.student.profile.language or statement.student.profile.language_adds:
		# 	if statement.student.profile.language:
		# 		p.add_run(str(statement.student.profile.language) + ' ').underline = True
		# 	p.add_run(statement.student.profile.language_adds).underline = True
		# else:
		# 	p.add_run('не изучал(а)').underline = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('Общежитие ')
		# if statement.motel:
		# 	p.add_run('нуждаюсь').underline = True
		# else:
		# 	p.add_run('не нуждаюсь').underline = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('О себе дополнительно сообщаю: ').bold = True
		# p.add_run('_______________________________________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('_____________________________________________________________________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('_____________________________________________________________________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('Ф.И.О. родителей, контактные телефоны (для очной формы обучения): ')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# if parents:
		# 	if parents[0].surname_patent_1 or parents[0].name_patent_1 or parents[0].name2_patent_1 or \
		# 	parents[0].surname_patent_2 or parents[0].name_patent_2 or parents[0].name2_patent_2:
		# 		p = document.add_paragraph()
		# 		if parents[0].surname_patent_1 or parents[0].name_patent_1 or parents[0].name2_patent_1:
		# 			p.add_run(parents[0].type1 + ': ' + parents[0].surname_patent_1 + ' ' + parents[0].name_patent_1 + ' ' + parents[0].name2_patent_1 + ' - ' + parents[0].phone_patent_1 + ' ').underline = True
		# 		if parents[0].surname_patent_2 or parents[0].name_patent_2 or parents[0].name2_patent_2:
		# 			p.add_run(parents[0].type2 + ': ' + parents[0].surname_patent_2 + ' ' + parents[0].name_patent_2 + ' ' + parents[0].name2_patent_2 + ' - ' + parents[0].phone_patent_2).underline = True
		# 		p.paragraph_format.space_after = 0
		# 		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('_____________________________________________________________________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# p = document.add_paragraph()
		# p.add_run('_____________________________________________________________________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

		# document.add_paragraph().paragraph_format.space_after = 0
		# document.add_paragraph().paragraph_format.space_after = 0


		# table = document.add_table(rows=1, cols=2)
		# table.allow_autifit = False
		# table.style = 'TableGrid'
		# table.columns[0].width = Mm(130)
		# table.columns[1].width = Mm(50)

		# hdr_cells = table.rows[0].cells
		# hdr_cells[0].text = 'Вопрос/ответ'
		# for paragraph in hdr_cells[0].paragraphs:
		# 	for run in paragraph.runs:
		# 		run.font.bold = True
		# hdr_cells[0].width = Mm(130)
		# hdr_cells[1].text = 'Подпись поступающего'
		# for paragraph in hdr_cells[1].paragraphs:
		# 	for run in paragraph.runs:
		# 		run.font.bold = True
		# hdr_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Среднее профессиональное образование по программе подготовки специалистов среднего звена получаю '
		# if statement.first_flag:
		# 	row_cells[0].text += 'ВПЕРВЫЕ'
		# else:
		# 	row_cells[0].text += 'НЕ ВПЕРВЫЕ'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Создание специальных условий при проведении вступительных испытаний в связи с инвалидностью \
		# или ограниченными возможностями здоровья '
		# if statement.special_flag:
		# 	row_cells[0].text += 'НУЖДАЮСЬ'
		# else:
		# 	row_cells[0].text += 'НЕ НУЖДАЮСЬ'
		# row_cells[0].text += '\nЕсли нуждаюсь, указать реквизиты документа, подтверждающего создание специальных условий (справка МСЭ или ПМПК) '
		# row_cells[0].text += '___________'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Документы, подтверждающие результаты индивидуальных достижений (п. 8.3 Правил приема) '
		# if statement.docs_flag:
		# 	row_cells[0].text += 'ИМЕЮ'
		# else:
		# 	row_cells[0].text += 'НЕ ИМЕЮ'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Договор о целевом обучении в данном образовательном учреждении '
		# if statement.docs_goals_flag:
		# 	row_cells[0].text += 'ИМЕЮ'
		# else:
		# 	row_cells[0].text += 'НЕ ИМЕЮ'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'Ознакомлен:'
		# for paragraph in row_cells[0].paragraphs:
		# 	for run in paragraph.runs:
		# 		run.font.bold = True
		# row_cells[0].width = Mm(130)
		# row_cells[1].text = 'Подпись поступающего'
		# for paragraph in row_cells[1].paragraphs:
		# 	for run in paragraph.runs:
		# 		run.font.bold = True
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'С Уставом, лицензией на право осуществления образовательной деятельности, свидетельством о \
		# государственной аккредитации и приложениями к ним, правилами приема и другими документами, регламентирующими \
		# организацию и осуществление образовательной деятельности, правами и обязанностями обучающихся в данном \
		# образовательном учреждении ОЗНАКОМЛЕН(А)'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'С правилами подачи апелляций при приеме на первый курс по результатам проведения вступительных \
		# испытаний, проводимых образовательным учреждением самостоятельно, ОЗНАКОМЛЕН(А)'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# row_cells = table.add_row().cells
		# row_cells[0].text = 'С датой предоставления оригинала документа об образовании и (или) документа об образовании и \
		# квалификаци, ОЗНАКОМЛЕН(А)'
		# row_cells[0].width = Mm(130)
		# row_cells[1].width = Mm(50)

		# document.add_paragraph().paragraph_format.space_after = 0
		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('Дата: "')
		# p.add_run(statement.send_date.strftime('%d')).underline = True
		# p.add_run('" ')
		# p.add_run(months[statement.send_date.month - 1]).underline = True
		# p.add_run(' ')
		# p.add_run(statement.send_date.strftime('%Y') + ' г.').underline = True
		# p.add_run('            _____________________________/_________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('                                                           (подпись поступающего)                   (расшифровка)').italic = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# document.add_paragraph().paragraph_format.space_after = 0
		# document.add_paragraph().paragraph_format.space_after = 0
		# document.add_paragraph().paragraph_format.space_after = 0
		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('Ответственное лицо')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('приемной комиссии                    _____________________________/_________________________')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# p = document.add_paragraph()
		# p.add_run('                                                                         (подпись)                                   (расшифровка)').italic = True
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT

		# document.add_paragraph().paragraph_format.space_after = 0

		# p = document.add_paragraph()
		# p.add_run('"___" ____________ 20__ г.')
		# p.paragraph_format.space_after = 0
		# p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.LEFT


		file_name = 'Statement-art'
		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=' + file_name +' (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response

	for picture in pictures:
		copictures_list[picture.pk] = CoPicturee.objects.filter(picture = picture)

	args = {
		'pictures': pictures,
		'copictures_list': copictures_list
	}
	return render(request, 'pictures/view_arts.html', args)


@login_required(login_url='/login/')
def load_image(request):
	if not request.user.profile.member_access:
		return redirect('home')

	author = request.user
	coprofiles = CoProfile.objects.filter(main_user = author, profile_type__in = ['1', '2', '3', '4', '5'])

	if request.method=='POST':
		
		form_img = PictureUploadForm(request.POST, request.FILES, author = author, label_suffix='')

		if form_img.is_valid():
			new_img = form_img.save(commit=False)
			new_img.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoPicturee.objects.filter(picture = new_img, coauthor = coprofile):
						CoPicturee.objects.create(picture = new_img, coauthor = coprofile)
				else:
					CoPicturee.objects.filter(picture = new_img, coauthor = coprofile).delete()

			return redirect('pictures:view_arts')

		args = {
			'form': form_img,
			'coprofiles': coprofiles		}
		return render(request, 'pictures/load_image.html', args)	
	
	form_img = PictureUploadForm(author = author, label_suffix='')

	args = {
		'form': form_img,
		'coprofiles': coprofiles	}
	return render(request, 'pictures/load_image.html', args)



@login_required(login_url='/login/')
def load_image_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	author = Profile.objects.get(pk = pk).user
	coprofiles = CoProfile.objects.filter(main_user = author)

	if request.method=='POST':
		
		if author.profile.participation == '2':
			form_img = PictureUploadForm(request.POST, request.FILES, label_suffix='')
		else:
			form_img = PictureUploadNoneFileForm(request.POST, label_suffix='')

		if form_img.is_valid():
			new_img = form_img.save(commit=False)
			new_img.author = author
			new_img.save()	
			
			if author.profile.participation == '2':
				if 'file' in request.FILES:
					new_img.file = request.FILES['file']
					new_img.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoPicturee.objects.filter(picture = new_img, coauthor = coprofile):
						CoPicturee.objects.create(picture = new_img, coauthor = coprofile)
				else:
					CoPicturee.objects.filter(picture = new_img, coauthor = coprofile).delete()

			return redirect('view_contestant', pk = author.profile.pk)

		args = {
			'form': form_img,
			'coprofiles': coprofiles,
			'author': author
		}
		return render(request, 'pictures/load_image.html', args)	
	
	if author.profile.participation == '2':
		form_img = PictureUploadForm(label_suffix='')
	else:
		form_img = PictureUploadNoneFileForm(label_suffix='')


	args = {
		'form': form_img,
		'coprofiles': coprofiles,
		'author': author
	}
	return render(request, 'pictures/load_image.html', args)


@login_required(login_url='/login/')
def edit_image(request, pk):
	if not request.user.profile.member_access:
		return redirect('home')
		
	pict = Picture.objects.get(pk=pk)
	author = request.user
	coprofiles = CoProfile.objects.filter(main_user = author, profile_type__in = ['1', '2', '3', '4', '5'])
	copictures_list = {}

	if pict.author.main_user != author:
		return redirect('home')

	if request.method=='POST':
		form_img = PictureEditForm(request.POST, request.FILES, instance=pict, author = author, label_suffix='')

		if form_img.is_valid():
			new_img = form_img.save(commit=False)
			new_img.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoPicturee.objects.filter(picture = new_img, coauthor = coprofile):
						CoPicturee.objects.create(picture = new_img, coauthor = coprofile)
				else:
					CoPicturee.objects.filter(picture = new_img, coauthor = coprofile).delete()


			return redirect('pictures:view_arts')

		args = {
			'form': form_img,
		}
		return render(request, 'pictures/edit_image.html', args)	
	
	form_img = PictureEditForm(instance=pict, author = author, label_suffix='')
	copicture = list(CoPicturee.objects.filter(picture = pict).values_list('coauthor', flat=True))
	for coprofile in coprofiles:
		copictures_list[coprofile.pk] = coprofile.pk in copicture

	args = {
		'form': form_img,
		'pict': pict,
		'coprofiles': coprofiles,
		'copictures_list': copictures_list
	}
	return render(request, 'pictures/edit_image.html', args)



@login_required(login_url='/login/')
def edit_image_admin(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')
		
	pict = Picture.objects.get(pk=pk)
	author = pict.author
	coprofiles = CoProfile.objects.filter(main_user = author)
	copictures_list = {}

	if pict.author != author:
		return redirect('home')

	if request.method=='POST':
		form_img = PictureEditForm(request.POST, request.FILES, instance=pict, label_suffix='')

		if form_img.is_valid():
			new_img = form_img.save(commit=False)
			new_img.author = author
			new_img.save()

			for coprofile in coprofiles:
				test_str = 'coprofile-check-' + str(coprofile.pk)
				if test_str in request.POST:
					if not CoPicturee.objects.filter(picture = new_img, coauthor = coprofile):
						CoPicturee.objects.create(picture = new_img, coauthor = coprofile)
				else:
					CoPicturee.objects.filter(picture = new_img, coauthor = coprofile).delete()


			return redirect('view_contestant', pk = author.profile.pk)

		args = {
			'form': form_img,
		}
		return render(request, 'pictures/edit_image.html', args)	
	
	form_img = PictureEditForm(instance=pict, label_suffix='')
	copicture = list(CoPicturee.objects.filter(picture = pict).values_list('coauthor', flat=True))
	for coprofile in coprofiles:
		copictures_list[coprofile.pk] = coprofile.pk in copicture

	args = {
		'form': form_img,
		'pict': pict,
		'coprofiles': coprofiles,
		'copictures_list': copictures_list
	}
	return render(request, 'pictures/edit_image.html', args)


# --------------------------------
#           Для ajax'а
# --------------------------------
@login_required(login_url='/login/')
def ajax_del_image(request):
	image_pk = request.GET['image']
	picture = Picture.objects.get(pk=image_pk)

	if not request.user.profile.admin_access and picture.author.main_user != request.user:
		return HttpResponse(False)

	picture.delete()

	return HttpResponse(True)
