import os
import datetime
import operator
import locale

from datetime import date

from django.conf import settings
from django.core.files.storage import default_storage

from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib.auth import update_session_auth_hash

from django.contrib.auth.models import User

from pytils import translit

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Mm, Pt

from pictures.models import Picture, CoPicturee
from movies.models import Movie, CoMovie
from nominations.models import ArtNomination, VocalNomination
from marks.models import PictureMark, MovieMark

from profileuser.models import CoProfile

from marks.forms import PictureMarkForm, MovieMarkForm


@login_required(login_url='/login/')
def view_art_nomination(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')

	nomination = ArtNomination.objects.get(pk=pk)

	pictures_all = Picture.objects.filter(nomination=nomination)
	copictures = {}
	for picture_all in pictures_all:
		copictures[picture_all.pk] = CoPicturee.objects.filter(picture=picture_all)

	pictures_1_1 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='1', participation='1')
	pictures_2_1 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='2', participation='1')
	pictures_3_1 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='3', participation='1')
	pictures_4_1 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='4', participation='1')
	pictures_5_1 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='5', participation='1')
	pictures_6_1 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='6', participation='1')

	pictures_1_2 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='1', participation='2')
	pictures_2_2 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='2', participation='2')
	pictures_3_2 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='3', participation='2')
	pictures_4_2 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='4', participation='2')
	pictures_5_2 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='5', participation='2')
	pictures_6_2 = Picture.objects.filter(nomination=nomination, author__main_user__profile__category='6', participation='2')

	marks_1_1 = PictureMark.objects.filter(work__in = pictures_1_1)
	marks_2_1 = PictureMark.objects.filter(work__in = pictures_2_1)
	marks_3_1 = PictureMark.objects.filter(work__in = pictures_3_1)
	marks_4_1 = PictureMark.objects.filter(work__in = pictures_4_1)
	marks_5_1 = PictureMark.objects.filter(work__in = pictures_5_1)
	marks_6_1 = PictureMark.objects.filter(work__in = pictures_6_1)

	marks_1_2 = PictureMark.objects.filter(work__in = pictures_1_2)
	marks_2_2 = PictureMark.objects.filter(work__in = pictures_2_2)
	marks_3_2 = PictureMark.objects.filter(work__in = pictures_3_2)
	marks_4_2 = PictureMark.objects.filter(work__in = pictures_4_2)
	marks_5_2 = PictureMark.objects.filter(work__in = pictures_5_2)
	marks_6_2 = PictureMark.objects.filter(work__in = pictures_6_2)

	ratings_1_1 = {}
	ratings_2_1 = {}
	ratings_3_1 = {}
	ratings_4_1 = {}
	ratings_5_1 = {}
	ratings_6_1 = {}

	ratings_1_2 = {}
	ratings_2_2 = {}
	ratings_3_2 = {}
	ratings_4_2 = {}
	ratings_5_2 = {}
	ratings_6_2 = {}

	pictures_list_1_1 = {}
	pictures_list_2_1 = {}
	pictures_list_3_1 = {}
	pictures_list_4_1 = {}
	pictures_list_5_1 = {}
	pictures_list_6_1 = {}

	pictures_list_1_2 = {}
	pictures_list_2_2 = {}
	pictures_list_3_2 = {}
	pictures_list_4_2 = {}
	pictures_list_5_2 = {}
	pictures_list_6_2 = {}


	for picture in pictures_1_1:
		mrks = marks_1_1.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_1_1[picture.id] = round(summa, 1)
		else:
			ratings_1_1[picture.id] = 0
		pictures_list_1_1[picture.id]=picture

	for picture in pictures_2_1:
		mrks = marks_2_1.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_2_1[picture.id] = round(summa, 1)
		else:
			ratings_2_1[picture.id] = 0
		pictures_list_2_1[picture.id]=picture

	for picture in pictures_3_1:
		mrks = marks_3_1.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_3_1[picture.id] = round(summa, 1)
		else:
			ratings_3_1[picture.id] = 0
		pictures_list_3_1[picture.id]=picture

	for picture in pictures_4_1:
		mrks = marks_4_1.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_4_1[picture.id] = round(summa, 1)
		else:
			ratings_4_1[picture.id] = 0
		pictures_list_4_1[picture.id]=picture

	for picture in pictures_5_1:
		mrks = marks_5_1.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_5_1[picture.id] = round(summa, 1)
		else:
			ratings_5_1[picture.id] = 0
		pictures_list_5_1[picture.id]=picture

	for picture in pictures_6_1:
		mrks = marks_6_1.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_6_1[picture.id] = round(summa, 1)
		else:
			ratings_6_1[picture.id] = 0
		pictures_list_6_1[picture.id]=picture


########################################################
	for picture in pictures_1_2:
		mrks = marks_1_2.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_1_2[picture.id] = round(summa, 1)
		else:
			ratings_1_2[picture.id] = 0
		pictures_list_1_2[picture.id]=picture

	for picture in pictures_2_2:
		mrks = marks_2_2.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_2_2[picture.id] = round(summa, 1)
		else:
			ratings_2_2[picture.id] = 0
		pictures_list_2_2[picture.id]=picture

	for picture in pictures_3_2:
		mrks = marks_3_2.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_3_2[picture.id] = round(summa, 1)
		else:
			ratings_3_2[picture.id] = 0
		pictures_list_3_2[picture.id]=picture

	for picture in pictures_4_2:
		mrks = marks_4_2.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_4_2[picture.id] = round(summa, 1)
		else:
			ratings_4_2[picture.id] = 0
		pictures_list_4_2[picture.id]=picture

	for picture in pictures_5_2:
		mrks = marks_5_2.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_5_2[picture.id] = round(summa, 1)
		else:
			ratings_5_2[picture.id] = 0
		pictures_list_5_2[picture.id]=picture

	for picture in pictures_6_2:
		mrks = marks_6_2.filter(work=picture)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
				summa += mark.criterai_four
				summa += mark.criterai_five
			summa = summa / length
			ratings_6_2[picture.id] = round(summa, 1)
		else:
			ratings_6_2[picture.id] = 0
		pictures_list_6_2[picture.id]=picture

	sorting_1_1 = sorted(ratings_1_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_2_1 = sorted(ratings_2_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_3_1 = sorted(ratings_3_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_4_1 = sorted(ratings_4_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_5_1 = sorted(ratings_5_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_6_1 = sorted(ratings_6_1.items(), key=operator.itemgetter(1), reverse=True)

	sorting_1_2 = sorted(ratings_1_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_2_2 = sorted(ratings_2_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_3_2 = sorted(ratings_3_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_4_2 = sorted(ratings_4_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_5_2 = sorted(ratings_5_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_6_2 = sorted(ratings_6_2.items(), key=operator.itemgetter(1), reverse=True)


	if request.POST:
		filename = translit.slugify(str(nomination)) + ' ('
		if 'type_1_1' in request.POST:
			pictures_list = pictures_list_1_1
			sorting = sorting_1_1
			name = 'Студенты (профи) высших учебных заведений (очное участие)'
			filename += 'students_vpo_profi_ochno (DPI)'
		elif 'type_1_2' in request.POST:
			pictures_list = pictures_list_1_2
			sorting = sorting_1_2
			name = 'Студенты (профи) высших учебных заведений (заочное участие)'
			filename += 'students_vpo_profi__zaochno (DPI)'
		elif  'type_2_1' in request.POST:
			pictures_list = pictures_list_2_1
			sorting = sorting_2_1
			name = 'Студенты (любители) высших учебных заведений (очное участие)'
			filename += 'students_vpo_lubiteli_ochno (DPI)'
		elif  'type_2_2' in request.POST:
			pictures_list = pictures_list_2_2
			sorting = sorting_2_2
			name = 'Студенты (любители) высших учебных заведений (заочное участие)'
			filename += 'students_vpo_lubiteli_zaochno (DPI)'
		elif  'type_3_1' in request.POST:
			pictures_list = pictures_list_3_1
			sorting = sorting_3_1
			name = 'Студенты (профи) учреждений среднего профессионального образовани (очное участие)'
			filename += 'students_spo_profi_ochno (DPI)'
		elif  'type_3_2' in request.POST:
			pictures_list = pictures_list_3_2
			sorting = sorting_3_2
			name = 'Студенты (профи) учреждений среднего профессионального образовани (заочное участие)'
			filename += 'students_spo_profi_zaochno (DPI)'
		elif  'type_4_1' in request.POST:
			pictures_list = pictures_list_4_1
			sorting = sorting_4_1
			name = 'Студенты (любители) учреждений среднего профессионального образовани (очное участие)'
			filename += 'students_spo_lubiteli_ochno (DPI)'
		elif  'type_4_2' in request.POST:
			pictures_list = pictures_list_4_2
			sorting = sorting_4_2
			name = 'Студенты (любители) учреждений среднего профессионального образовани (заочное участие)'
			filename += 'students_spo_lubiteli_zaochno (DPI)'
		elif  'type_5_1' in request.POST:
			pictures_list = pictures_list_5_1
			sorting = sorting_5_1
			name = 'Профи (очное участие)'
			filename += 'profi_ochno (DPI)'
		elif  'type_5_2' in request.POST:
			pictures_list = pictures_list_5_2
			sorting = sorting_5_2
			name = 'Профи (заочное участие)'
			filename += 'profi_zaochno (DPI)'
		elif  'type_6_1' in request.POST:
			pictures_list = pictures_list_6_1
			sorting = sorting_6_1
			name = 'Любители (очное участие)'
			filename += 'lubiteli_ochno (DPI)'
		elif  'type_6_2' in request.POST:
			pictures_list = pictures_list_6_2
			sorting = sorting_6_2
			name = 'Любители (заочное участие)'
			filename += 'lubiteli_zaochno (DPI)'
		filename += ')'


		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()
		document = Document()
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = Mm(297)
		section.page_height = Mm(210)
		section.left_margin = Mm(30)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)


		document.add_paragraph(str(nomination) + ' (' + name + ')').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		p = document.add_paragraph()
		p.add_run(dte.strftime('%d %B %Y')).italic = True
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

		table = document.add_table(rows=1, cols=6)
		table.allow_autifit = False
		table.style = 'TableGrid'
		table.columns[0].width = Mm(10)
		table.columns[1].width = Mm(60)
		table.columns[2].width = Mm(60)
		table.columns[3].width = Mm(50)
		table.columns[4].width = Mm(50)
		table.columns[5].width = Mm(17)

		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = '№'
		hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[0].width = Mm(10)
		hdr_cells[1].text = 'Название работы'
		hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[1].width = Mm(60)
		hdr_cells[2].text = 'Конкурсант'
		hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[2].width = Mm(60)
		hdr_cells[3].text = 'Преподаватель'
		hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[3].width = Mm(50)
		hdr_cells[4].text = 'Учреждение'
		hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[4].width = Mm(50)
		hdr_cells[5].text = 'Баллы'
		hdr_cells[5].paragraphs[0].runs[0].font.bold = True
		hdr_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[5].width = Mm(17)


		cnt = 1
		for key, val in sorting:
			picture = pictures_list[key]

			row_cells = table.add_row().cells
			row_cells[0].text = str(cnt)
			row_cells[0].paragraphs[0].runs[0].font.bold = True
			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(10)
			row_cells[1].text = picture.name
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(60)
			row_cells[2].text = picture.author.get_full_name()
			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[2].width = Mm(60)
			row_cells[3].text = ''
			for copicture in copictures[key]:
				row_cells[3].text += copicture.coauthor.get_file_name() + ' (' + copicture.coauthor.get_profile_type_display() + ')\n'
			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[3].width = Mm(50)
			row_cells[4].text = picture.author.main_user.profile.institution
			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[4].width = Mm(50)
			row_cells[5].text = str(val)
			row_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[5].width = Mm(17)

			cnt += 1

		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=' + filename +' (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response

	args = {
		'nomination': nomination, 
		'pictures_list_1_1': pictures_list_1_1,
		'pictures_list_2_1': pictures_list_2_1,
		'pictures_list_3_1': pictures_list_3_1,
		'pictures_list_4_1': pictures_list_4_1,
		'pictures_list_5_1': pictures_list_5_1,
		'pictures_list_6_1': pictures_list_6_1,
		'pictures_list_1_2': pictures_list_1_2,
		'pictures_list_2_2': pictures_list_2_2,
		'pictures_list_3_2': pictures_list_3_2,
		'pictures_list_4_2': pictures_list_4_2,
		'pictures_list_5_2': pictures_list_5_2,
		'pictures_list_6_2': pictures_list_6_2,
		'sorting_1_1': sorting_1_1,
		'sorting_2_1': sorting_2_1,
		'sorting_3_1': sorting_3_1,
		'sorting_4_1': sorting_4_1,
		'sorting_5_1': sorting_5_1,
		'sorting_6_1': sorting_6_1,
		'sorting_1_2': sorting_1_2,
		'sorting_2_2': sorting_2_2,
		'sorting_3_2': sorting_3_2,
		'sorting_4_2': sorting_4_2,
		'sorting_5_2': sorting_5_2,
		'sorting_6_2': sorting_6_2,
		'copictures': copictures
	}
	return render(request, 'ratings/view_art_nominations.html', args)


@login_required(login_url='/login/')
def view_mov_nomination(request, pk):
	if not request.user.profile.admin_access:
		return redirect('home')


	nomination = VocalNomination.objects.get(pk=pk)

	movies_all = Movie.objects.filter(nomination=nomination)
	comovies = {}
	for move_all in movies_all:
		comovies[move_all.pk] = CoMovie.objects.filter(movie=move_all)

	movies_1_1 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='1', participation='1')
	movies_2_1 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='2', participation='1')
	movies_3_1 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='3', participation='1')
	movies_4_1 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='4', participation='1')
	movies_5_1 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='5', participation='1')
	movies_6_1 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='6', participation='1')

	movies_1_2 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='1', participation='2')
	movies_2_2 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='2', participation='2')
	movies_3_2 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='3', participation='2')
	movies_4_2 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='4', participation='2')
	movies_5_2 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='5', participation='2')
	movies_6_2 = Movie.objects.filter(nomination=nomination, author__main_user__profile__category='6', participation='2')

	marks_1_1 = MovieMark.objects.filter(work__in = movies_1_1)
	marks_2_1 = MovieMark.objects.filter(work__in = movies_2_1)
	marks_3_1 = MovieMark.objects.filter(work__in = movies_3_1)
	marks_4_1 = MovieMark.objects.filter(work__in = movies_4_1)
	marks_5_1 = MovieMark.objects.filter(work__in = movies_5_1)
	marks_6_1 = MovieMark.objects.filter(work__in = movies_6_1)

	marks_1_2 = MovieMark.objects.filter(work__in = movies_1_2)
	marks_2_2 = MovieMark.objects.filter(work__in = movies_2_2)
	marks_3_2 = MovieMark.objects.filter(work__in = movies_3_2)
	marks_4_2 = MovieMark.objects.filter(work__in = movies_4_2)
	marks_5_2 = MovieMark.objects.filter(work__in = movies_5_2)
	marks_6_2 = MovieMark.objects.filter(work__in = movies_6_2)

	ratings_1_1 = {}
	ratings_2_1 = {}
	ratings_3_1 = {}
	ratings_4_1 = {}
	ratings_5_1 = {}
	ratings_6_1 = {}

	ratings_1_2 = {}
	ratings_2_2 = {}
	ratings_3_2 = {}
	ratings_4_2 = {}
	ratings_5_2 = {}
	ratings_6_2 = {}

	movies_list_1_1 = {}
	movies_list_2_1 = {}
	movies_list_3_1 = {}
	movies_list_4_1 = {}
	movies_list_5_1 = {}
	movies_list_6_1 = {}

	movies_list_1_2 = {}
	movies_list_2_2 = {}
	movies_list_3_2 = {}
	movies_list_4_2 = {}
	movies_list_5_2 = {}
	movies_list_6_2 = {}



	for movie in movies_1_1:
		mrks = marks_1_1.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_1_1[movie.id] = round(summa, 1)
		else:
			ratings_1_1[movie.id] = 0
		movies_list_1_1[movie.id]=movie

	for movie in movies_2_1:
		mrks = marks_2_1.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_2_1[movie.id] = round(summa, 1)
		else:
			ratings_2_1[movie.id] = 0
		movies_list_2_1[movie.id]=movie

	for movie in movies_3_1:
		mrks = marks_3_1.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_3_1[movie.id] = round(summa, 1)
		else:
			ratings_3_1[movie.id] = 0
		movies_list_3_1[movie.id]=movie

	for movie in movies_4_1:
		mrks = marks_4_1.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_4_1[movie.id] = round(summa, 1)
		else:
			ratings_4_1[movie.id] = 0
		movies_list_4_1[movie.id]=movie

	for movie in movies_5_1:
		mrks = marks_5_1.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_5_1[movie.id] = round(summa, 1)
		else:
			ratings_5_1[movie.id] = 0
		movies_list_5_1[movie.id]=movie

	for movie in movies_6_1:
		mrks = marks_6_1.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_6_1[movie.id] = round(summa, 1)
		else:
			ratings_6_1[movie.id] = 0
		movies_list_6_1[movie.id]=movie

########################################################
	for movie in movies_1_2:
		mrks = marks_1_2.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_1_2[movie.id] = round(summa, 1)
		else:
			ratings_1_2[movie.id] = 0
		movies_list_1_2[movie.id]=movie

	for movie in movies_2_2:
		mrks = marks_2_2.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_2_2[movie.id] = round(summa, 1)
		else:
			ratings_2_2[movie.id] = 0
		movies_list_2_2[movie.id]=movie

	for movie in movies_3_2:
		mrks = marks_3_2.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_3_2[movie.id] = round(summa, 1)
		else:
			ratings_3_2[movie.id] = 0
		movies_list_3_2[movie.id]=movie

	for movie in movies_4_2:
		mrks = marks_4_2.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_4_2[movie.id] = round(summa, 1)
		else:
			ratings_4_2[movie.id] = 0
		movies_list_4_2[movie.id]=movie

	for movie in movies_5_2:
		mrks = marks_5_2.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_5_2[movie.id] = round(summa, 1)
		else:
			ratings_5_2[movie.id] = 0
		movies_list_5_2[movie.id]=movie

	for movie in movies_6_2:
		mrks = marks_6_2.filter(work=movie)
		if mrks:
			length = len(mrks)
			summa = 0;
			for mark in mrks:
				summa += mark.criterai_one
				summa += mark.criterai_two
				summa += mark.criterai_three
			summa = summa / length
			ratings_6_2[movie.id] = round(summa, 1)
		else:
			ratings_6_2[movie.id] = 0
		movies_list_6_2[movie.id]=movie


	sorting_1_1 = sorted(ratings_1_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_2_1 = sorted(ratings_2_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_3_1 = sorted(ratings_3_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_4_1 = sorted(ratings_4_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_5_1 = sorted(ratings_5_1.items(), key=operator.itemgetter(1), reverse=True)
	sorting_6_1 = sorted(ratings_6_1.items(), key=operator.itemgetter(1), reverse=True)

	sorting_1_2 = sorted(ratings_1_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_2_2 = sorted(ratings_2_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_3_2 = sorted(ratings_3_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_4_2 = sorted(ratings_4_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_5_2 = sorted(ratings_5_2.items(), key=operator.itemgetter(1), reverse=True)
	sorting_6_2 = sorted(ratings_6_2.items(), key=operator.itemgetter(1), reverse=True)

	if request.POST:
		filename = translit.slugify(str(nomination)) + ' ('
		if 'type_1_1' in request.POST:
			movies_list = movies_list_1_1
			sorting = sorting_1_1
			name = 'Студенты (профи) высших учебных заведений (очное участие)'
			filename += 'students_vpo_profi_ochno (vocal)'
		elif 'type_1_2' in request.POST:
			movies_list = movies_list_1_2
			sorting = sorting_1_2
			name = 'Студенты (профи) высших учебных заведений (заочное участие)'
			filename += 'students_vpo_profi__zaochno (vocal)'
		elif  'type_2_1' in request.POST:
			movies_list = movies_list_2_1
			sorting = sorting_2_1
			name = 'Студенты (любители) высших учебных заведений (очное участие)'
			filename += 'students_vpo_lubiteli_ochno (vocal)'
		elif  'type_2_2' in request.POST:
			movies_list = movies_list_2_2
			sorting = sorting_2_2
			name = 'Студенты (любители) высших учебных заведений (заочное участие)'
			filename += 'students_vpo_lubiteli_zaochno (vocal)'
		elif  'type_3_1' in request.POST:
			movies_list = movies_list_3_1
			sorting = sorting_3_1
			name = 'Студенты (профи) учреждений среднего профессионального образовани (очное участие)'
			filename += 'students_spo_profi_ochno (vocal)'
		elif  'type_3_2' in request.POST:
			movies_list = movies_list_3_2
			sorting = sorting_3_2
			name = 'Студенты (профи) учреждений среднего профессионального образовани (заочное участие)'
			filename += 'students_spo_profi_zaochno (vocal)'
		elif  'type_4_1' in request.POST:
			movies_list = movies_list_4_1
			sorting = sorting_4_1
			name = 'Студенты (любители) учреждений среднего профессионального образовани (очное участие)'
			filename += 'students_spo_lubiteli_ochno (vocal)'
		elif  'type_4_2' in request.POST:
			movies_list = movies_list_4_2
			sorting = sorting_4_2
			name = 'Студенты (любители) учреждений среднего профессионального образовани (заочное участие)'
			filename += 'students_spo_lubiteli_zaochno (vocal)'
		elif  'type_5_1' in request.POST:
			movies_list = movies_list_5_1
			sorting = sorting_5_1
			name = 'Профи (очное участие)'
			filename += 'profi_ochno (vocal)'
		elif  'type_5_2' in request.POST:
			movies_list = movies_list_5_2
			sorting = sorting_5_2
			name = 'Профи (заочное участие)'
			filename += 'profi_zaochno (vocal)'
		elif  'type_6_1' in request.POST:
			movies_list = movies_list_6_1
			sorting = sorting_6_1
			name = 'Любители (очное участие)'
			filename += 'lubiteli_ochno (vocal)'
		elif  'type_6_2' in request.POST:
			movies_list = movies_list_6_2
			sorting = sorting_6_2
			name = 'Любители (заочное участие)'
			filename += 'lubiteli_zaochno (vocal)'
		filename += ')'

		locale.setlocale(locale.LC_ALL, ('ru_RU', 'UTF-8'))

		dte = date.today()
		document = Document()
		section = document.sections[-1]
		new_width, new_height = section.page_height, section.page_width
		section.orientation = WD_ORIENT.LANDSCAPE
		section.page_width = Mm(297)
		section.page_height = Mm(210)
		section.left_margin = Mm(30)
		section.right_margin = Mm(10)
		section.top_margin = Mm(10)
		section.bottom_margin = Mm(10)
		section.header_distance = Mm(10)
		section.footer_distance = Mm(10)

		style = document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(12)


		document.add_paragraph(str(nomination) + ' (' + name + ')').paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		p = document.add_paragraph()
		p.add_run(dte.strftime('%d %B %Y')).italic = True
		p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

		table = document.add_table(rows=1, cols=6)
		table.allow_autifit = False
		table.style = 'TableGrid'
		table.columns[0].width = Mm(10)
		table.columns[1].width = Mm(60)
		table.columns[2].width = Mm(60)
		table.columns[3].width = Mm(50)
		table.columns[4].width = Mm(50)
		table.columns[5].width = Mm(17)

		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = '№'
		hdr_cells[0].paragraphs[0].runs[0].font.bold = True
		hdr_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[0].width = Mm(10)
		hdr_cells[1].text = 'Название работы'
		hdr_cells[1].paragraphs[0].runs[0].font.bold = True
		hdr_cells[1].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[1].width = Mm(60)
		hdr_cells[2].text = 'Конкурсант(Коллектив)'
		hdr_cells[2].paragraphs[0].runs[0].font.bold = True
		hdr_cells[2].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[2].width = Mm(60)
		hdr_cells[3].text = 'Преподаватель/Концертмейстер'
		hdr_cells[3].paragraphs[0].runs[0].font.bold = True
		hdr_cells[3].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[3].width = Mm(50)
		hdr_cells[4].text = 'Учреждение'
		hdr_cells[4].paragraphs[0].runs[0].font.bold = True
		hdr_cells[4].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[4].width = Mm(50)
		hdr_cells[5].text = 'Баллы'
		hdr_cells[5].paragraphs[0].runs[0].font.bold = True
		hdr_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
		hdr_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		hdr_cells[5].width = Mm(17)


		cnt = 1
		for key, val in sorting:
			movie = movies_list[key]

			row_cells = table.add_row().cells
			row_cells[0].text = str(cnt)
			row_cells[0].paragraphs[0].runs[0].font.bold = True
			row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[0].width = Mm(10)
			row_cells[1].text = movie.name_1
			if movie.composer_1:
				row_cells[1].text += ' муз. ' + movie.composer_1
			if movie.poet_1:
				row_cells[1].text += ' сл. ' + movie.poet_1
			if movie.region_1:
				row_cells[1].text += '\nРегион: ' + movie.region_1
			row_cells[1].text += '\n' + movie.name_2
			if movie.composer_2:
				row_cells[1].text += ' муз. ' + movie.composer_2
			if movie.poet_2:
				row_cells[1].text += ' сл. ' + movie.poet_2
			if movie.region_2:
				row_cells[1].text += '\nРегион: ' + movie.region_2
			row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[1].width = Mm(60)
			row_cells[2].text = movie.author.get_full_name()
			row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[2].width = Mm(60)
			row_cells[3].text = ''
			for comove in comovies[key]:
				row_cells[3].text += comove.coauthor.get_file_name() + ' (' + comove.coauthor.get_profile_type_display() + ')\n'
			row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[3].width = Mm(50)
			row_cells[4].text = movie.author.main_user.profile.institution
			row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[4].width = Mm(50)
			row_cells[5].text = str(val)
			row_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
			row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
			row_cells[5].width = Mm(17)

			cnt += 1

		response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
		response['Content-Disposition'] = 'attachment; filename=' + filename +' (' + dte.strftime('%d-%m-%Y') + ').docx'
		document.save(response)

		return response

	args = {
		'nomination': nomination, 
		'movies_list_1_1': movies_list_1_1,
		'movies_list_2_1': movies_list_2_1,
		'movies_list_3_1': movies_list_3_1,
		'movies_list_4_1': movies_list_4_1,
		'movies_list_5_1': movies_list_5_1,
		'movies_list_6_1': movies_list_6_1,
		'movies_list_1_2': movies_list_1_2,
		'movies_list_2_2': movies_list_2_2,
		'movies_list_3_2': movies_list_3_2,
		'movies_list_4_2': movies_list_4_2,
		'movies_list_5_2': movies_list_5_2,
		'movies_list_6_2': movies_list_6_2,
		'sorting_1_1': sorting_1_1,
		'sorting_2_1': sorting_2_1,
		'sorting_3_1': sorting_3_1,
		'sorting_4_1': sorting_4_1,
		'sorting_5_1': sorting_5_1,
		'sorting_6_1': sorting_6_1,
		'sorting_1_2': sorting_1_2,
		'sorting_2_2': sorting_2_2,
		'sorting_3_2': sorting_3_2,
		'sorting_4_2': sorting_4_2,
		'sorting_5_2': sorting_5_2,
		'sorting_6_2': sorting_6_2,
		'comovies': comovies
	}
	return render(request, 'ratings/view_mov_nominations.html', args)



@login_required(login_url='/login/')
def view_protocols(request):
	if not request.user.profile.admin_access:
		return redirect('home')

	juries = User.objects.filter(profile__juri_accecc = True)

	args = {
		'juries': juries, 
	}
	return render(request, 'ratings/view_protocols.html', args)


@login_required(login_url='/login/')
def get_check_list(request, pk, param):
	if not request.user.profile.admin_access:
		return redirect('home')


	CATEGORY_TYPES = {
		'1': 'Студенты (профи) высших учебных заведений',
		'2': 'Студенты (любители) высших учебных заведений',
		'3': 'Студенты (профи) учреждений среднего профессионального образовани',
		'4': 'Студенты (любители) учреждений среднего профессионального образовани',
		'5': 'Профи',
		'6': 'Любители',
	}

	user = User.objects.get(pk=pk)
	juri_type = user.profile.juri_type

	document = Document()
	section = document.sections[-1]
	new_width, new_height = section.page_height, section.page_width
	section.orientation = WD_ORIENT.LANDSCAPE
	section.page_width = Mm(297)
	section.page_height = Mm(210)
	section.left_margin = Mm(10)
	section.right_margin = Mm(10)
	section.top_margin = Mm(10)
	section.bottom_margin = Mm(10)
	section.header_distance = Mm(10)
	section.footer_distance = Mm(10)

	style = document.styles['Normal']
	font = style.font
	font.name = 'Times New Roman'
	font.size = Pt(8)

	p = document.add_paragraph()
	p.add_run('Оценочный лист Всероссийского фестиваля конкурса народного творчества «ГАВРИЛОВСКИЕ ГУЛЯНИЯ», 2023г.').bold = True
	p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
	p.paragraph_format.space_after = 0
	document.add_paragraph()

	if juri_type == '1':
		#Вокал
		criteria1 = 'Сложность и трактовка музыкальных произведений'
		criteria2 = 'Интонационная выразительность'
		criteria3 = 'Артистизм'


		filename = translit.slugify(str(user.profile.get_file_name()) + ' (оценочный лист)')

		nominations = VocalNomination.objects.all()

		for nomination in nominations:

			find_cnt = Movie.objects.filter(nomination=nomination, participation = param).count()
			if find_cnt:
				document.add_paragraph()
				p = document.add_paragraph()
				p.add_run(nomination.name).bold = True
				p.paragraph_format.space_after = 0

				for cat_num in ['1','2','3','4','5','6']:
					category = CATEGORY_TYPES[cat_num]


					members = Movie.objects.filter(nomination=nomination, author__main_user__profile__category=cat_num, participation = param)

					if members:
						p = document.add_paragraph()
						p.paragraph_format.space_after = 0
						p = document.add_paragraph(category)
						p.paragraph_format.space_after = 0
		
						table = document.add_table(rows=2, cols=9)
						table.allow_autifit = False
						table.style = 'TableGrid'
						table.columns[0].width = Mm(10)
						table.columns[1].width = Mm(50)
						table.columns[2].width = Mm(50)
						table.columns[3].width = Mm(50)
						table.columns[4].width = Mm(40)


						table.columns[5].width = Mm(20)
						table.columns[6].width = Mm(20)
						table.columns[7].width = Mm(20)

						table.columns[8].width = Mm(17)

						new_cell = table.cell(0, 0).merge(table.cell(1, 0))
						new_cell.text = '№'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 1).merge(table.cell(1, 1))
						new_cell.text = 'Участник'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 2).merge(table.cell(1, 2))
						new_cell.text = 'Учреждение'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 3).merge(table.cell(1, 3))
						new_cell.text = 'Преподаватель'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 4).merge(table.cell(1, 4))
						new_cell.text = 'Название'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						
						new_cell = table.cell(0, 5).merge(table.cell(0, 7))
						new_cell.text = 'Оценка (макс. 10 баллов)'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						table.cell(1, 5).text = criteria1
						table.cell(1, 6).text = criteria2
						table.cell(1, 7).text = criteria3

						new_cell = table.cell(0, 8).merge(table.cell(1, 8))
						new_cell.text = 'Всего'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						cnt = 1
						for member in members:
							marks = MovieMark.objects.filter(work=member, expert=user)

							row_cells = table.add_row().cells
							row_cells[0].text = str(cnt)
							row_cells[0].paragraphs[0].runs[0].font.bold = True
							row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[0].width = Mm(10)

							row_cells[1].text = member.author.get_file_name()
							row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[1].width = Mm(50)


							row_cells[2].text = member.author.main_user.profile.institution
							row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[2].width = Mm(50)

							row_cells[3].text = ''
							teachers = CoMovie.objects.filter(movie = member)
							if teachers:
								for teacher in teachers:
									row_cells[3].text += teacher.coauthor.get_profile_type_display() + ' ' + teacher.coauthor.get_file_name() + '\n'
							row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[3].width = Mm(50)

							row_cells[4].text = member.name_1
							if member.composer_1:
								row_cells[4].text += ' муз. ' + member.composer_1
							if member.poet_1:
								row_cells[4].text += ' сл. ' + member.poet_1
							if member.region_1:
								row_cells[4].text += '\nРегион: ' + member.region_1
							row_cells[4].text += '\n' + member.name_2
							if member.composer_2:
								row_cells[4].text += ' муз. ' + member.composer_2
							if member.poet_2:
								row_cells[4].text += ' сл. ' + member.poet_2
							if member.region_2:
								row_cells[4].text += '\nРегион: ' + member.region_2
							row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[4].width = Mm(40)

							
							row_cells[5].text = ''
							if marks:
								row_cells[5].text = str(marks[0].criterai_one)
							row_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[5].width = Mm(20)

							row_cells[6].text = ''
							if marks:
								row_cells[6].text = str(marks[0].criterai_two)
							row_cells[6].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[6].width = Mm(20)

							row_cells[7].text = ''
							if marks:
								row_cells[7].text = str(marks[0].criterai_three)
							row_cells[7].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[7].width = Mm(20)


							row_cells[8].text = ''
							row_cells[8].width = Mm(17)

							summa = 0
							if marks:
								if marks[0].criterai_one:
									summa += marks[0].criterai_one
								if marks[0].criterai_two:
									summa += marks[0].criterai_two
								if marks[0].criterai_three:
									summa += marks[0].criterai_three
							row_cells[8].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[8].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[8].text = str(summa)

							cnt += 1


	if juri_type == '2':
		#Вокал
		criteria1 = 'Соответствие названию, полнота раскрытия'
		criteria2 = 'Техническое воспроизведение'
		criteria3 = 'Авторское новаторство'
		criteria4 = 'Эстетика подачи работы'
		criteria5 = 'Визуальное восприятие'


		filename = translit.slugify(str(user.profile.get_file_name()) + ' (оценочный лист)')

		nominations = ArtNomination.objects.all()

		for nomination in nominations:

			find_cnt = Picture.objects.filter(nomination=nomination, participation = param).count()
			if find_cnt:
				document.add_paragraph()
				p = document.add_paragraph()
				p.add_run(nomination.name).bold = True
				p.paragraph_format.space_after = 0

				for cat_num in ['1','2','3','4']:
					category = CATEGORY_TYPES[cat_num]


					members = Picture.objects.filter(nomination=nomination, author__main_user__profile__category=cat_num, participation = param)

					if members:
						p = document.add_paragraph()
						p.paragraph_format.space_after = 0
						p = document.add_paragraph(category)
						p.paragraph_format.space_after = 0
		
						table = document.add_table(rows=2, cols=11)
						table.allow_autifit = False
						table.style = 'TableGrid'
						table.columns[0].width = Mm(10)
						table.columns[1].width = Mm(40)
						table.columns[2].width = Mm(40)
						table.columns[3].width = Mm(40)
						table.columns[4].width = Mm(30)


						table.columns[5].width = Mm(20)
						table.columns[6].width = Mm(20)
						table.columns[7].width = Mm(20)
						table.columns[8].width = Mm(20)
						table.columns[9].width = Mm(20)

						table.columns[10].width = Mm(17)

						new_cell = table.cell(0, 0).merge(table.cell(1, 0))
						new_cell.text = '№'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 1).merge(table.cell(1, 1))
						new_cell.text = 'Участник'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 2).merge(table.cell(1, 2))
						new_cell.text = 'Учреждение'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 3).merge(table.cell(1, 3))
						new_cell.text = 'Преподаватель'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
						new_cell = table.cell(0, 4).merge(table.cell(1, 4))
						new_cell.text = 'Название'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						
						new_cell = table.cell(0, 5).merge(table.cell(0, 9))
						new_cell.text = 'Оценка (макс. 10 баллов)'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						table.cell(1, 5).text = criteria1
						table.cell(1, 6).text = criteria2
						table.cell(1, 7).text = criteria3
						table.cell(1, 8).text = criteria4
						table.cell(1, 9).text = criteria5

						new_cell = table.cell(0, 10).merge(table.cell(1, 10))
						new_cell.text = 'Всего'
						new_cell.paragraphs[0].runs[0].font.bold = True
						new_cell.paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
						new_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

						cnt = 1
						for member in members:
							marks = PictureMark.objects.filter(work=member, expert=user)

							row_cells = table.add_row().cells
							row_cells[0].text = str(cnt)
							row_cells[0].paragraphs[0].runs[0].font.bold = True
							row_cells[0].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[0].width = Mm(10)

							row_cells[1].text = member.author.get_file_name()
							row_cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[1].width = Mm(40)


							row_cells[2].text = member.author.main_user.profile.institution
							row_cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[2].width = Mm(40)

							row_cells[3].text = ''
							teachers = CoPicturee.objects.filter(picture = member)
							if teachers:
								for teacher in teachers:
									row_cells[3].text += teacher.coauthor.get_profile_type_display() + ' ' + teacher.coauthor.get_file_name() + '\n'
							row_cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[3].width = Mm(40)

							row_cells[4].text = member.name
							row_cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[4].width = Mm(30)

							
							row_cells[5].text = ''
							if marks:
								row_cells[5].text = str(marks[0].criterai_one)
							row_cells[5].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[5].width = Mm(20)

							row_cells[6].text = ''
							if marks:
								row_cells[6].text = str(marks[0].criterai_two)
							row_cells[6].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[6].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[6].width = Mm(20)

							row_cells[7].text = ''
							if marks:
								row_cells[7].text = str(marks[0].criterai_three)
							row_cells[7].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[7].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[7].width = Mm(20)

							row_cells[8].text = ''
							if marks:
								row_cells[8].text = str(marks[0].criterai_four)
							row_cells[8].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[8].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[8].width = Mm(20)

							row_cells[9].text = ''
							if marks:
								row_cells[9].text = str(marks[0].criterai_five)
							row_cells[9].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[9].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[9].width = Mm(20)


							row_cells[10].text = ''
							row_cells[10].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[10].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
							row_cells[10].width = Mm(17)

							summa = 0
							if marks:
								if marks[0].criterai_one:
									summa += marks[0].criterai_one
								if marks[0].criterai_two:
									summa += marks[0].criterai_two
								if marks[0].criterai_three:
									summa += marks[0].criterai_three
								if marks[0].criterai_four:
									summa += marks[0].criterai_four
								if marks[0].criterai_five:
									summa += marks[0].criterai_five
							row_cells[10].text = str(summa)
							row_cells[10].paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
							row_cells[10].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

							cnt += 1

	document.add_paragraph()
	p = document.add_paragraph()
	p.add_run('______________________ /'  + user.profile.get_file_name() + '/').bold = True
	p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.RIGHT

	response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
	response['Content-Disposition'] = 'attachment; filename=' + filename + '.docx'
	document.save(response)

	return response



